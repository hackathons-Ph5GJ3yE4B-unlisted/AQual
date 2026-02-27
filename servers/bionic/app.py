from flask import Flask, request, jsonify, render_template, Response
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import nsmap
import re
import io
import html
import json
import base64
import asyncio
import sys
import time
import uuid
import urllib.parse
import hashlib
import logging
from threading import Lock
from pathlib import Path
from collections import deque
import httpx
from google import genai
from google.genai import types
from lxml import html as lxml_html

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from servers.config import get_env

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.logger.setLevel("INFO")
SUPPRESS_RING_POLL_REQUEST_LOGS = str(get_env("SUPPRESS_RING_POLL_REQUEST_LOGS", "1")).strip().lower() not in ("0", "false", "no", "off")
RING_POLL_DEBUG = str(get_env("RING_POLL_DEBUG", "0")).strip().lower() in ("1", "true", "yes", "on")
READING_MODE_DEBUG = str(get_env("READING_MODE_DEBUG", "1")).strip().lower() in ("1", "true", "yes", "on")


class _SuppressRingPollAccessLogs(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        try:
            message = record.getMessage()
        except Exception:
            return True
        return "/ring-event/poll" not in message


if SUPPRESS_RING_POLL_REQUEST_LOGS:
    werkzeug_logger = logging.getLogger("werkzeug")
    werkzeug_logger.addFilter(_SuppressRingPollAccessLogs())
    # Suppress access-log spam from long-poll endpoints while keeping app-level logs.
    werkzeug_logger.setLevel(logging.ERROR)

GEMINI_API_KEY = get_env("GEMINI_API_KEY", "")
GEMINI_MODEL = "gemini-3-flash-preview"
GEMINI_THINKING_LEVEL = "LOW"
GEMINI_READING_MODE_MODEL = str(get_env("GEMINI_READING_MODE_MODEL", "")).strip() or "gemini-3-flash-preview"
GEMINI_READING_MODE_THINKING_LEVEL = str(get_env("GEMINI_READING_MODE_THINKING_LEVEL", "HIGH")).strip().upper() or "HIGH"
if GEMINI_READING_MODE_THINKING_LEVEL not in ("LOW", "MEDIUM", "HIGH"):
    GEMINI_READING_MODE_THINKING_LEVEL = "HIGH"
GEMINI_READING_MODE_USE_THINKING = str(get_env("GEMINI_READING_MODE_USE_THINKING", "0")).strip().lower() in ("1", "true", "yes", "on")
GEMINI_READING_MODE_ENABLE_AI_REFINEMENT = str(get_env("GEMINI_READING_MODE_ENABLE_AI_REFINEMENT", "0")).strip().lower() in ("1", "true", "yes", "on")
GEMINI_LIVE_MODEL = str(get_env("GEMINI_LIVE_MODEL", "")).strip() or "gemini-2.5-flash-native-audio-preview-12-2025"
GEMINI_LIVE_MODEL_FALLBACKS = [
    str(get_env("GEMINI_LIVE_FALLBACK_MODEL_1", "")).strip() or "gemini-2.5-flash-native-audio-preview-09-2025",
    str(get_env("GEMINI_LIVE_FALLBACK_MODEL_2", "")).strip(),
]
ELEVENLABS_API_KEY = get_env("ELEVENLABS_API_KEY", "")
ELEVENLABS_TTS_VOICE_ID = get_env("ELEVENLABS_TTS_VOICE_ID", "21m00Tcm4TlvDq8ikWAM")
ELEVENLABS_TTS_MODEL_ID = get_env("ELEVENLABS_TTS_MODEL_ID", "eleven_multilingual_v2")


def _get_env_int(name: str, default: int, minimum=None, maximum=None) -> int:
    raw = str(get_env(name, default)).strip()
    try:
        value = int(raw)
    except Exception:
        value = int(default)
    if minimum is not None:
        value = max(int(minimum), value)
    if maximum is not None:
        value = min(int(maximum), value)
    return value


def _get_env_float(name: str, default: float, minimum=None, maximum=None) -> float:
    raw = str(get_env(name, default)).strip()
    try:
        value = float(raw)
    except Exception:
        value = float(default)
    if minimum is not None:
        value = max(float(minimum), value)
    if maximum is not None:
        value = min(float(maximum), value)
    return value


GEMINI_LIVE_THINKING_BUDGET = _get_env_int("GEMINI_LIVE_THINKING_BUDGET", 0, minimum=0, maximum=32768)
GEMINI_LIVE_MAX_OUTPUT_TOKENS = _get_env_int("GEMINI_LIVE_MAX_OUTPUT_TOKENS", 160, minimum=32, maximum=2048)
GEMINI_LIVE_TEMPERATURE = _get_env_float("GEMINI_LIVE_TEMPERATURE", 0.1, minimum=0.0, maximum=2.0)
GEMINI_READING_MODE_HTML_CHAR_LIMIT = _get_env_int("GEMINI_READING_MODE_HTML_CHAR_LIMIT", 450000, minimum=20000, maximum=1200000)
GEMINI_READING_MODE_MAX_OUTPUT_TOKENS = _get_env_int("GEMINI_READING_MODE_MAX_OUTPUT_TOKENS", 900, minimum=128, maximum=4096)
READING_MODE_DISABLE_SELECTOR_CAP = str(get_env("READING_MODE_DISABLE_SELECTOR_CAP", "0")).strip().lower() in ("1", "true", "yes", "on")
READING_MODE_URL_CACHE_ENABLED = str(get_env("READING_MODE_URL_CACHE_ENABLED", "1")).strip().lower() in ("1", "true", "yes", "on")
READING_MODE_URL_CACHE_MAX_ENTRIES = _get_env_int("READING_MODE_URL_CACHE_MAX_ENTRIES", 5000, minimum=100, maximum=50000)
READING_MODE_URL_CACHE_FILE = PROJECT_ROOT / ".run" / "reading_mode_url_cache.json"


def get_gemini_client():
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not configured. Set it in .env.")
    return genai.Client(
        api_key=GEMINI_API_KEY,
        http_options={"api_version": "v1beta"},
    )

# Store images temporarily for the session
image_store = {}
LIVE_SESSION_TTL_SECONDS = 2 * 60 * 60
LIVE_SESSION_MAX_ENTRIES = 200
live_session_store = {}
live_session_lock = Lock()
live_screenshot_store = {}
ring_event_lock = Lock()
ring_event_counter = 0
ring_event_last_ts = 0.0
ring_event_history = deque(maxlen=500)
READING_MODE_PLAN_CACHE_TTL_SECONDS = 90.0
READING_MODE_PLAN_CACHE_MAX_ENTRIES = 120
if READING_MODE_DISABLE_SELECTOR_CAP:
    READING_MODE_MAX_INCLUDE_SELECTORS = 0
    READING_MODE_MAX_EXCLUDE_SELECTORS = 0
else:
    READING_MODE_MAX_INCLUDE_SELECTORS = _get_env_int("READING_MODE_MAX_INCLUDE_SELECTORS", 8, minimum=1, maximum=32)
    READING_MODE_MAX_EXCLUDE_SELECTORS = _get_env_int("READING_MODE_MAX_EXCLUDE_SELECTORS", 140, minimum=1, maximum=1000)
READING_MODE_PLAN_CACHE_VERSION = "2026-02-27-v5"
reading_mode_plan_cache = {}
reading_mode_plan_cache_lock = Lock()
reading_mode_url_cache = {}
reading_mode_url_cache_loaded = False
reading_mode_url_cache_lock = Lock()


def _set_cors_headers(response):
    origin = request.headers.get("Origin")
    response.headers["Access-Control-Allow-Origin"] = origin or "*"
    response.headers["Vary"] = "Origin"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    if request.headers.get("Access-Control-Request-Private-Network") == "true":
        response.headers["Access-Control-Allow-Private-Network"] = "true"
    return response


@app.after_request
def add_cors_headers(response):
    return _set_cors_headers(response)


def _log_gemini_event(event: str, **fields):
    payload = {"event": event, **fields}
    app.logger.info("[gemini] %s", json.dumps(payload, ensure_ascii=True, separators=(",", ":")))


def _log_reading_mode_event(event: str, **fields):
    if not READING_MODE_DEBUG:
        return
    payload = {"event": event, **fields}
    app.logger.info("[reading-mode] %s", json.dumps(payload, ensure_ascii=True, separators=(",", ":")))


def _log_tts_event(event: str, **fields):
    payload = {"event": event, **fields}
    app.logger.info("[tts] %s", json.dumps(payload, ensure_ascii=True, separators=(",", ":")))


def _log_ring_event(event: str, **fields):
    payload = {"event": event, **fields}
    app.logger.info("[ring] %s", json.dumps(payload, ensure_ascii=True, separators=(",", ":")))


def _register_ring_event(source: str = "unknown", payload=None):
    global ring_event_counter, ring_event_last_ts
    now_ts = time.time()
    payload_data = payload if isinstance(payload, dict) else {}
    with ring_event_lock:
        ring_event_counter += 1
        ring_event_last_ts = now_ts
        cursor = ring_event_counter
        ring_event_history.append({
            "cursor": int(cursor),
            "timestamp": float(now_ts),
            "source": str(source or "unknown"),
            "payload": payload_data,
        })
    _log_ring_event(
        "event_push",
        source=source,
        cursor=cursor,
        timestamp=now_ts,
        payload=payload or {},
    )
    return cursor, now_ts


def _extract_host(url: str) -> str:
    url = (url or "").strip()
    if not url:
        return ""
    try:
        return urllib.parse.urlparse(url).netloc or ""
    except Exception:
        return ""


def _normalize_live_conversation_id(raw_value) -> str:
    token = re.sub(r"[^a-zA-Z0-9._:-]+", "", str(raw_value or ""))
    return token[:120]


def _normalize_live_model_name(model_name: str) -> str:
    token = str(model_name or "").strip()
    if not token:
        return token
    if "/" in token:
        return token
    return f"models/{token}"


def _get_screenshot_hash(image_bytes: bytes) -> str:
    if not image_bytes:
        return ""
    return hashlib.sha1(image_bytes).hexdigest()[:16]


def _should_send_screenshot(conversation_id: str, screenshot_hash: str) -> bool:
    if not screenshot_hash:
        return False
    if not conversation_id:
        return True
    now_ts = time.time()
    with live_session_lock:
        _prune_live_session_store(now_ts)
        previous = str(live_screenshot_store.get(conversation_id) or "")
        live_screenshot_store[conversation_id] = screenshot_hash
        return previous != screenshot_hash


def _prune_live_session_store(now_ts: float):
    expired = [key for key, value in live_session_store.items() if now_ts - float(value.get("updated_at", 0)) > LIVE_SESSION_TTL_SECONDS]
    for key in expired:
        live_session_store.pop(key, None)
        live_screenshot_store.pop(key, None)
    if len(live_session_store) <= LIVE_SESSION_MAX_ENTRIES:
        return
    ordered = sorted(live_session_store.items(), key=lambda item: float(item[1].get("updated_at", 0)))
    for key, _value in ordered[: max(0, len(live_session_store) - LIVE_SESSION_MAX_ENTRIES)]:
        live_session_store.pop(key, None)
        live_screenshot_store.pop(key, None)
    for key in list(live_screenshot_store.keys()):
        if key not in live_session_store:
            live_screenshot_store.pop(key, None)


def _get_live_session_handle(conversation_id: str) -> str:
    if not conversation_id:
        return ""
    now_ts = time.time()
    with live_session_lock:
        _prune_live_session_store(now_ts)
        entry = live_session_store.get(conversation_id)
        if not entry:
            return ""
        if now_ts - float(entry.get("updated_at", 0)) > LIVE_SESSION_TTL_SECONDS:
            live_session_store.pop(conversation_id, None)
            return ""
        entry["updated_at"] = now_ts
        return str(entry.get("handle") or "")


def _set_live_session_handle(conversation_id: str, handle: str):
    if not conversation_id:
        return
    handle = str(handle or "").strip()
    if not handle:
        return
    now_ts = time.time()
    with live_session_lock:
        _prune_live_session_store(now_ts)
        live_session_store[conversation_id] = {
            "handle": handle,
            "updated_at": now_ts,
        }


def _prune_reading_mode_plan_cache(now_ts: float):
    stale_keys = [
        cache_key
        for cache_key, entry in reading_mode_plan_cache.items()
        if now_ts - float(entry.get("updated_at", 0.0)) > READING_MODE_PLAN_CACHE_TTL_SECONDS
    ]
    for cache_key in stale_keys:
        reading_mode_plan_cache.pop(cache_key, None)

    while len(reading_mode_plan_cache) > READING_MODE_PLAN_CACHE_MAX_ENTRIES:
        oldest_key = min(
            reading_mode_plan_cache,
            key=lambda cache_key: float(reading_mode_plan_cache[cache_key].get("updated_at", 0.0)),
        )
        reading_mode_plan_cache.pop(oldest_key, None)


def _get_cached_reading_mode_plan(cache_key: str):
    token = str(cache_key or "").strip()
    if not token:
        return None
    now_ts = time.time()
    with reading_mode_plan_cache_lock:
        _prune_reading_mode_plan_cache(now_ts)
        entry = reading_mode_plan_cache.get(token)
        if not entry:
            return None
        payload = dict(entry.get("payload") or {})
        if str(payload.get("planVersion") or "") != READING_MODE_PLAN_CACHE_VERSION:
            return None
        return payload


def _set_cached_reading_mode_plan(cache_key: str, payload: dict):
    token = str(cache_key or "").strip()
    if not token or not isinstance(payload, dict):
        return
    now_ts = time.time()
    with reading_mode_plan_cache_lock:
        _prune_reading_mode_plan_cache(now_ts)
        reading_mode_plan_cache[token] = {
            "payload": dict(payload),
            "updated_at": now_ts,
        }


def _normalize_reading_mode_cache_url(page_url: str) -> str:
    raw = str(page_url or "").strip()
    if not raw:
        return ""
    try:
        parsed = urllib.parse.urlsplit(raw)
    except Exception:
        return raw[:2000]
    scheme = str(parsed.scheme or "").lower()
    if scheme not in ("http", "https"):
        return ""
    host = str(parsed.netloc or "").lower()
    path = str(parsed.path or "/")
    query = str(parsed.query or "")
    normalized = urllib.parse.urlunsplit((scheme, host, path, query, ""))
    return normalized[:2000]


def _prune_reading_mode_url_cache_locked():
    while len(reading_mode_url_cache) > READING_MODE_URL_CACHE_MAX_ENTRIES:
        oldest_key = min(
            reading_mode_url_cache,
            key=lambda cache_key: float(reading_mode_url_cache[cache_key].get("updated_at", 0.0)),
        )
        reading_mode_url_cache.pop(oldest_key, None)


def _load_reading_mode_url_cache_locked():
    global reading_mode_url_cache_loaded
    if reading_mode_url_cache_loaded:
        return
    reading_mode_url_cache_loaded = True

    path = READING_MODE_URL_CACHE_FILE
    if not path.exists():
        return
    try:
        payload = json.loads(path.read_text("utf-8"))
    except Exception:
        return
    if not isinstance(payload, dict):
        return

    reading_mode_url_cache.clear()
    for key, value in payload.items():
        cache_url = _normalize_reading_mode_cache_url(key)
        if not cache_url or not isinstance(value, dict):
            continue
        cached_payload = value.get("payload")
        if not isinstance(cached_payload, dict):
            continue
        updated_at = float(value.get("updated_at", 0.0) or 0.0)
        reading_mode_url_cache[cache_url] = {
            "payload": dict(cached_payload),
            "updated_at": updated_at or time.time(),
        }
    _prune_reading_mode_url_cache_locked()


def _persist_reading_mode_url_cache_locked():
    try:
        READING_MODE_URL_CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            key: {
                "payload": value.get("payload") if isinstance(value, dict) else {},
                "updated_at": float(value.get("updated_at", 0.0) or 0.0) if isinstance(value, dict) else 0.0,
            }
            for key, value in reading_mode_url_cache.items()
        }
        READING_MODE_URL_CACHE_FILE.write_text(
            json.dumps(payload, ensure_ascii=True, separators=(",", ":")),
            encoding="utf-8",
        )
    except Exception:
        # URL cache persistence is best-effort.
        return


def _get_cached_reading_mode_plan_by_url(page_url: str):
    if not READING_MODE_URL_CACHE_ENABLED:
        return None
    cache_url = _normalize_reading_mode_cache_url(page_url)
    if not cache_url:
        return None
    with reading_mode_url_cache_lock:
        _load_reading_mode_url_cache_locked()
        entry = reading_mode_url_cache.get(cache_url)
        if not isinstance(entry, dict):
            return None
        payload = entry.get("payload")
        if not isinstance(payload, dict) or not payload:
            return None
        if str(payload.get("planVersion") or "") != READING_MODE_PLAN_CACHE_VERSION:
            return None
        entry["updated_at"] = time.time()
        return dict(payload)


def _set_cached_reading_mode_plan_by_url(page_url: str, payload: dict):
    if not READING_MODE_URL_CACHE_ENABLED:
        return
    cache_url = _normalize_reading_mode_cache_url(page_url)
    if not cache_url or not isinstance(payload, dict) or not payload:
        return
    with reading_mode_url_cache_lock:
        _load_reading_mode_url_cache_locked()
        reading_mode_url_cache[cache_url] = {
            "payload": dict(payload),
            "updated_at": time.time(),
        }
        _prune_reading_mode_url_cache_locked()
        _persist_reading_mode_url_cache_locked()


def _clean_ai_json_text(response_text: str) -> str:
    response_text = (response_text or "").strip()
    if response_text.startswith("```"):
        lines = response_text.split("\n")
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()
    return response_text


def _extract_first_json_object_text(response_text: str) -> str:
    text = str(response_text or "")
    start = text.find("{")
    if start < 0:
        return ""

    depth = 0
    in_string = False
    escape = False
    for index in range(start, len(text)):
        ch = text[index]
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == "\"":
                in_string = False
            continue
        if ch == "\"":
            in_string = True
            continue
        if ch == "{":
            depth += 1
            continue
        if ch == "}":
            depth -= 1
            if depth == 0:
                return text[start:index + 1]
    return ""


def _normalize_selector_list(raw_value, *, max_items: int):
    values = raw_value if isinstance(raw_value, (list, tuple)) else []
    normalized = []
    seen = set()
    cap_enabled = isinstance(max_items, int) and max_items > 0
    for item in values:
        token = str(item or "").strip()
        if not token:
            continue
        if len(token) > 220:
            continue
        lowered = token.lower()
        if lowered in ("*", "html", "body", ":root"):
            continue
        if any(fragment in lowered for fragment in ("script", "<", ">", "javascript:", "\n", "\r")):
            continue
        if lowered in seen:
            continue
        seen.add(lowered)
        normalized.append(token)
        if cap_enabled and len(normalized) >= max_items:
            break
    return normalized


def _is_aggressive_exclude_selector(selector: str) -> bool:
    token = str(selector or "").strip().lower()
    if not token:
        return True

    # Extremely broad primitives that can wipe most readable content.
    broad_primitives = {
        "div", "section", "article", "main", "span", "p", "a",
        "ul", "ol", "li", "table", "figure", "img", "video",
        "h1", "h2", "h3", "h4", "h5", "h6",
    }
    if token in broad_primitives:
        return True
    if token in ("*", "body *", "main *", "article *", "[role='main'] *", "[role=\"main\"] *"):
        return True
    if token.startswith("*"):
        return True
    if any(fragment in token for fragment in (":not(", ":has(", ":is(", ":where(")):
        return True

    # If selector lacks any targeting structure and no noise hint, it is usually too broad.
    has_structure = any(fragment in token for fragment in ("#", ".", "[", " ", ">", "+", "~", ":"))
    has_noise_hint = any(
        hint in token
        for hint in (
            "ad", "promo", "related", "recommend", "tag", "chip",
            "newsletter", "share", "social", "cookie", "consent",
            "sponsor", "sidebar", "rail", "card",
        )
    )
    if not has_structure and not has_noise_hint:
        return True
    return False


def _normalize_ai_exclude_selector_list(raw_value, *, max_items: int):
    double_cap = max_items * 2 if isinstance(max_items, int) and max_items > 0 else 0
    normalized = _normalize_selector_list(raw_value, max_items=double_cap)
    filtered = []
    seen = set()
    cap_enabled = isinstance(max_items, int) and max_items > 0
    for selector in normalized:
        lowered = str(selector or "").strip().lower()
        if not lowered or lowered in seen:
            continue
        seen.add(lowered)
        if _is_aggressive_exclude_selector(selector):
            continue
        filtered.append(selector)
        if cap_enabled and len(filtered) >= max_items:
            break
    return filtered


READING_MODE_POSITIVE_HINTS = (
    "article",
    "story",
    "content",
    "main",
    "post",
    "entry",
    "body",
    "text",
    "headline",
    "copy",
)

READING_MODE_NEGATIVE_HINTS = (
    "nav",
    "menu",
    "header",
    "footer",
    "cookie",
    "consent",
    "banner",
    "promo",
    "advert",
    "sponsor",
    "related",
    "recommend",
    "newsletter",
    "share",
    "social",
    "comment",
    "toolbar",
    "sidebar",
    "rail",
    "outbrain",
    "taboola",
    "most-read",
    "breaking-news",
    "topic",
    "topics",
    "tag",
    "tags",
    "keyword",
    "keywords",
    "taxonomy",
    "chip",
    "chips",
    "cluster",
    "meta",
    "metadata",
    "topic-list",
    "topiclist",
    "tag-list",
    "tags-list",
    "read-more",
    "most-read",
    "most-popular",
    "you-may-like",
    "recommended",
    "suggested",
    "inline-promo",
)

READING_MODE_BASELINE_EXCLUDE_SELECTORS = (
    "nav",
    "header",
    "footer",
    "aside",
    "[aria-label*='advert' i]",
    "[class*='ad-' i]",
    "[id*='ad-' i]",
    "[class*='advert' i]",
    "[id*='advert' i]",
    "[class*='cookie' i]",
    "[id*='cookie' i]",
    "[class*='consent' i]",
    "[id*='consent' i]",
    "[class*='newsletter' i]",
    "[id*='newsletter' i]",
    "[class*='related' i]",
    "[id*='related' i]",
    "[class*='recommend' i]",
    "[id*='recommend' i]",
    "[class*='share' i]",
    "[id*='share' i]",
    "[class*='social' i]",
    "[id*='social' i]",
    "[class*='topic' i]",
    "[id*='topic' i]",
    "[class*='topic-list' i]",
    "[id*='topic-list' i]",
    "[class*='topiclist' i]",
    "[id*='topiclist' i]",
    "[class*='tag' i]",
    "[id*='tag' i]",
    "[class*='tag-list' i]",
    "[id*='tag-list' i]",
    "[class*='tags-list' i]",
    "[id*='tags-list' i]",
    "[class*='keyword' i]",
    "[id*='keyword' i]",
    "[class*='taxonomy' i]",
    "[id*='taxonomy' i]",
    "[class*='chip' i]",
    "[id*='chip' i]",
    "[class*='cluster' i]",
    "[id*='cluster' i]",
    "[class*='clusteritems' i]",
    "[id*='clusteritems' i]",
    "[class*='cluster-items' i]",
    "[id*='cluster-items' i]",
    "[class*='card' i]",
    "[id*='card' i]",
    "[class*='promo' i]",
    "[id*='promo' i]",
    "[class*='read-more' i]",
    "[id*='read-more' i]",
    "[class*='most-read' i]",
    "[id*='most-read' i]",
    "[class*='most-popular' i]",
    "[id*='most-popular' i]",
    "[data-component*='topic' i]",
    "[data-component*='tag' i]",
    "[data-component*='related' i]",
    "[data-component*='promo' i]",
    "[data-testid*='topic' i]",
    "[data-testid*='tag' i]",
    "[data-testid*='related' i]",
    "[aria-label*='topic' i]",
    "[aria-label*='tag' i]",
    "[href*='/news/topics/' i]",
    "[href*='/topics/' i]",
    "[href*='/tags/' i]",
)

READING_MODE_CLUTTER_PRIORITY_HINTS = (
    "topic",
    "tag",
    "taxonomy",
    "keyword",
    "chip",
    "cluster",
    "related",
    "recommend",
    "promo",
    "card",
    "newsletter",
    "share",
    "social",
    "comment",
    "sponsor",
    "advert",
    "sidebar",
    "rail",
    "toolbar",
    "read-more",
    "most-read",
    "most-popular",
    "suggested",
    "you-may-like",
)


def _exclude_selector_priority(selector: str) -> int:
    token = str(selector or "").strip().lower()
    if not token:
        return -9999

    score = 0
    if any(hint in token for hint in ("topic", "tag", "taxonomy", "keyword", "chip", "cluster")):
        score += 2000
    if any(hint in token for hint in ("related", "recommend", "promo", "card", "read-more", "most-read", "most-popular", "suggested", "you-may-like")):
        score += 1600
    if any(hint in token for hint in ("share", "social", "comment", "newsletter")):
        score += 1300
    if any(hint in token for hint in ("advert", "sponsor", "outbrain", "taboola")) or "ad-" in token:
        score += 1100
    if any(hint in token for hint in ("sidebar", "rail", "toolbar")):
        score += 900
    if any(hint in token for hint in ("nav", "menu", "header", "footer")):
        score += 500
    if token.startswith("[class*=") or token.startswith("[id*=") or token.startswith("[data-") or token.startswith("[aria-"):
        score += 250
    if any(fragment in token for fragment in (" ", ">", "+", "~")):
        score += 70

    # Prefer concise reusable selectors over brittle long chains.
    score -= min(len(token), 220) // 5
    return score


def _sort_exclude_selectors_by_priority(selectors, *, max_items: int):
    triple_cap = max_items * 3 if isinstance(max_items, int) and max_items > 0 else 0
    normalized = _normalize_selector_list(selectors, max_items=triple_cap)
    unique = []
    seen = set()
    for selector in normalized:
        lowered = str(selector or "").strip().lower()
        if not lowered or lowered in seen:
            continue
        seen.add(lowered)
        unique.append(selector)
    unique.sort(
        key=lambda selector: (
            _exclude_selector_priority(selector),
            -len(str(selector or "")),
        ),
        reverse=True,
    )
    if isinstance(max_items, int) and max_items > 0:
        return unique[:max_items]
    return unique


def _merge_reading_mode_exclude_selectors(ai_exclude, heuristic_exclude, *, max_items: int):
    baseline = _normalize_selector_list(READING_MODE_BASELINE_EXCLUDE_SELECTORS, max_items=max_items)
    baseline_set = {str(selector or "").strip().lower() for selector in baseline}

    heuristic_dynamic = [
        selector
        for selector in _normalize_selector_list(heuristic_exclude, max_items=max_items * 2)
        if str(selector or "").strip().lower() not in baseline_set
    ]
    ai_dynamic = [
        selector
        for selector in _normalize_ai_exclude_selector_list(ai_exclude, max_items=max_items)
        if str(selector or "").strip().lower() not in baseline_set
    ]

    ordered = []
    ordered.extend(baseline)
    ordered.extend(_sort_exclude_selectors_by_priority(ai_dynamic, max_items=max_items))
    ordered.extend(_sort_exclude_selectors_by_priority(heuristic_dynamic, max_items=max_items))
    return _normalize_selector_list(ordered, max_items=max_items)


def _safe_css_simple_token(raw_value: str) -> str:
    token = str(raw_value or "").strip()
    if not token:
        return ""
    token = re.sub(r"[^A-Za-z0-9_-]", "", token)
    if not token:
        return ""
    if not re.match(r"^[A-Za-z_][A-Za-z0-9_-]*$", token):
        return ""
    return token


def _split_css_class_tokens(raw_value: str):
    raw_tokens = str(raw_value or "").strip().split()
    output = []
    seen = set()
    for raw_token in raw_tokens:
        token = _safe_css_simple_token(raw_token)
        if not token:
            continue
        lowered = token.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        output.append(token)
        if len(output) >= 3:
            break
    return output


def _build_selector_from_node(tag_name: str, node_id: str, node_class: str) -> str:
    tag = str(tag_name or "").strip().lower()
    if not re.match(r"^[a-z][a-z0-9]*$", tag):
        tag = "div"

    safe_id = _safe_css_simple_token(node_id)
    if safe_id:
        return f"#{safe_id}"

    class_tokens = _split_css_class_tokens(node_class)
    if class_tokens:
        return f"{tag}" + "".join(f".{token}" for token in class_tokens[:2])

    if tag in ("article", "main", "section"):
        return tag
    return ""


def _score_reading_mode_candidate(tag: str, id_class_text: str, text_length: int) -> int:
    score = int(text_length)
    lowered_tag = str(tag or "").strip().lower()
    lowered_text = str(id_class_text or "").strip().lower()

    if lowered_tag in ("article", "main"):
        score += 380
    elif lowered_tag == "section":
        score += 180
    elif lowered_tag == "div":
        score += 40

    if any(hint in lowered_text for hint in READING_MODE_POSITIVE_HINTS):
        score += 260
    if any(hint in lowered_text for hint in READING_MODE_NEGATIVE_HINTS):
        score -= 640
    return score


def _build_heuristic_reading_mode_plan(source_html: str):
    include_candidates = []
    exclude_candidates = []
    seen_excludes = set()

    try:
        parser = lxml_html.HTMLParser(encoding="utf-8", recover=True)
        root = lxml_html.fromstring(source_html, parser=parser)
    except Exception as e:
        return {
            "include_selectors": ["main article", "article", "main", "[role='main']", "#content", ".article"],
            "exclude_selectors": [
                "nav", "header", "footer", "aside",
                "[aria-label*='advert' i]", "[class*='ad-' i]", "[id*='ad-' i]",
                "[class*='cookie' i]", "[id*='cookie' i]",
                "[class*='consent' i]", "[id*='consent' i]",
                "[class*='newsletter' i]", "[id*='newsletter' i]",
                "[class*='related' i]", "[id*='related' i]",
                "[class*='share' i]", "[id*='share' i]",
            ],
            "candidate_include": [],
            "error": str(e)[:200],
        }

    for noisy in root.xpath("//script|//style|//noscript|//template|//svg|//canvas"):
        parent = noisy.getparent()
        if parent is not None:
            parent.remove(noisy)

    for node in root.iter():
        tag = str(getattr(node, "tag", "") or "").lower()
        if not tag or not re.match(r"^[a-z][a-z0-9]*$", tag):
            continue
        if tag in ("script", "style", "noscript", "template", "meta", "link"):
            continue

        node_id = str(node.get("id") or "")
        node_class = str(node.get("class") or "")
        id_class_text = f"{node_id} {node_class}".strip().lower()
        selector = _build_selector_from_node(tag, node_id, node_class)
        if not selector:
            continue

        text_content = " ".join(str(node.text_content() or "").split())
        text_length = len(text_content)
        if text_length < 120:
            if any(hint in id_class_text for hint in READING_MODE_NEGATIVE_HINTS):
                lowered = selector.lower()
                if lowered not in seen_excludes:
                    seen_excludes.add(lowered)
                    exclude_candidates.append({
                        "selector": selector,
                        "priority": _exclude_selector_priority(selector),
                        "text_length": text_length,
                    })
            continue

        score = _score_reading_mode_candidate(tag, id_class_text, text_length)
        include_candidates.append({
            "selector": selector,
            "tag": tag,
            "text_length": text_length,
            "score": score,
        })

        if any(hint in id_class_text for hint in READING_MODE_NEGATIVE_HINTS):
            lowered = selector.lower()
            if lowered not in seen_excludes:
                seen_excludes.add(lowered)
                exclude_candidates.append({
                    "selector": selector,
                    "priority": _exclude_selector_priority(selector),
                    "text_length": text_length,
                })

    include_candidates.sort(key=lambda item: (int(item.get("score", 0)), int(item.get("text_length", 0))), reverse=True)

    include_selectors = []
    include_seen = set()
    for candidate in include_candidates:
        selector = str(candidate.get("selector") or "").strip()
        if not selector:
            continue
        lowered = selector.lower()
        if lowered in include_seen:
            continue
        if any(hint in lowered for hint in ("nav", "menu", "header", "footer", "cookie", "consent", "share", "related")):
            continue
        if int(candidate.get("score", 0)) < 220:
            continue
        include_seen.add(lowered)
        include_selectors.append(selector)
        if len(include_selectors) >= 6:
            break

    if not include_selectors:
        include_selectors = ["main article", "article", "main", "[role='main']", "#content", ".article"]

    exclude_candidates.sort(
        key=lambda item: (
            int(item.get("priority", 0)),
            -int(item.get("text_length", 0)),
        ),
        reverse=True,
    )
    ordered_dynamic_excludes = [str(item.get("selector") or "").strip() for item in exclude_candidates]
    exclude_selectors = _normalize_selector_list(
        list(READING_MODE_BASELINE_EXCLUDE_SELECTORS) + ordered_dynamic_excludes,
        max_items=READING_MODE_MAX_EXCLUDE_SELECTORS,
    )

    return {
        "include_selectors": _normalize_selector_list(include_selectors, max_items=READING_MODE_MAX_INCLUDE_SELECTORS),
        "exclude_selectors": exclude_selectors,
        "candidate_include": include_candidates[:24],
        "error": "",
    }


def _extract_relaxed_selector_array(raw_text: str, key_name: str):
    text = str(raw_text or "")
    key_match = re.search(rf'"{re.escape(key_name)}"\s*:\s*\[', text, flags=re.IGNORECASE)
    if not key_match:
        return []

    cursor = key_match.end()
    in_string = False
    escape = False
    current = ""
    values = []
    while cursor < len(text):
        ch = text[cursor]
        if in_string:
            if escape:
                current += ch
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == "\"":
                token = current.strip()
                if token:
                    values.append(token)
                current = ""
                in_string = False
            elif ch in ("\n", "\r"):
                token = current.strip()
                if token:
                    values.append(token)
                current = ""
                break
            else:
                current += ch
        else:
            if ch == "]":
                break
            if ch == "\"":
                in_string = True
        cursor += 1

    if in_string and current.strip():
        values.append(current.strip())
    return values


def _parse_relaxed_reading_mode_payload(response_text: str):
    text = _clean_ai_json_text(response_text)
    include_values = _extract_relaxed_selector_array(text, "include_selectors")
    exclude_values = _extract_relaxed_selector_array(text, "exclude_selectors")
    notes_match = re.search(r'"notes"\s*:\s*"([^"\r\n]{0,400})', text, flags=re.IGNORECASE)
    notes = str(notes_match.group(1) if notes_match else "").strip()
    return {
        "include_selectors": _normalize_selector_list(include_values, max_items=READING_MODE_MAX_INCLUDE_SELECTORS),
        "exclude_selectors": _normalize_selector_list(exclude_values, max_items=READING_MODE_MAX_EXCLUDE_SELECTORS),
        "notes": notes[:400],
    }


def _decode_data_url(data_url: str):
    match = re.match(r"^data:(?P<mime>[^;,]+)?(?P<b64>;base64)?,(?P<data>.*)$", data_url, re.DOTALL)
    if not match:
        raise ValueError("Invalid data URL")

    mime_type = (match.group("mime") or "image/png").strip()
    payload = match.group("data") or ""
    if match.group("b64"):
        image_bytes = base64.b64decode(payload)
    else:
        image_bytes = urllib.parse.unquote_to_bytes(payload)
    return image_bytes, mime_type


def _load_web_image_payload(data: dict):
    image_data = data.get("imageData")
    content_type = (data.get("contentType") or "").split(";", 1)[0].strip()
    image_url = (data.get("imageUrl") or "").strip()

    if image_data:
        image_bytes = base64.b64decode(image_data)
        if not content_type:
            content_type = "image/jpeg"
    elif image_url:
        if image_url.startswith("data:"):
            image_bytes, content_type = _decode_data_url(image_url)
        else:
            response = httpx.get(
                image_url,
                follow_redirects=True,
                timeout=20.0,
                headers={"User-Agent": "AQual/1.0"},
            )
            response.raise_for_status()
            image_bytes = response.content
            content_type = (response.headers.get("content-type") or "image/jpeg").split(";", 1)[0].strip()
    else:
        raise ValueError("No image input provided")

    if not image_bytes:
        raise ValueError("Image is empty")
    if len(image_bytes) > 15 * 1024 * 1024:
        raise ValueError("Image is too large")
    if not content_type.startswith("image/"):
        content_type = "image/jpeg"

    return image_bytes, content_type


def _load_live_audio_payload(data: dict):
    audio_data = (data.get("audioData") or "").strip()
    audio_mime_type = (data.get("audioMimeType") or "audio/pcm;rate=24000").strip()
    if not audio_data:
        raise ValueError("No audio input provided")

    try:
        audio_bytes = base64.b64decode(audio_data, validate=True)
    except Exception:
        raise ValueError("Invalid base64 audio payload")

    if not audio_bytes:
        raise ValueError("Audio is empty")
    if len(audio_bytes) > 15 * 1024 * 1024:
        raise ValueError("Audio payload is too large")
    if not audio_mime_type:
        audio_mime_type = "audio/pcm;rate=24000"
    return audio_bytes, audio_mime_type


def _load_live_screenshot_payload(data: dict):
    screenshot_data_url = (data.get("screenshotDataUrl") or "").strip()
    screenshot_data = (data.get("screenshotData") or "").strip()
    screenshot_mime_type = (data.get("screenshotMimeType") or "image/png").split(";", 1)[0].strip()

    if screenshot_data_url:
        if not screenshot_data_url.startswith("data:"):
            raise ValueError("Screenshot must be provided as a data URL")
        screenshot_bytes, screenshot_mime_type = _decode_data_url(screenshot_data_url)
    elif screenshot_data:
        try:
            screenshot_bytes = base64.b64decode(screenshot_data, validate=True)
        except Exception:
            raise ValueError("Invalid base64 screenshot payload")
    else:
        raise ValueError("No screenshot provided")

    if not screenshot_bytes:
        raise ValueError("Screenshot is empty")
    if len(screenshot_bytes) > 15 * 1024 * 1024:
        raise ValueError("Screenshot is too large")
    if not screenshot_mime_type.startswith("image/"):
        screenshot_mime_type = "image/png"
    return screenshot_bytes, screenshot_mime_type


def _iter_live_models():
    ordered = [GEMINI_LIVE_MODEL] + list(GEMINI_LIVE_MODEL_FALLBACKS)
    seen = set()
    for model_name in ordered:
        token = str(model_name or "").strip()
        if not token or token in seen:
            continue
        seen.add(token)
        yield token


def _normalize_reading_mode_model_name(model_name: str) -> str:
    token = str(model_name or "").strip()
    if not token:
        return ""
    alias_map = {
        "gemini-3.0-flash": "gemini-3-flash-preview",
        "gemini-3-flash": "gemini-3-flash-preview",
    }
    return alias_map.get(token.lower(), token)


def _iter_reading_mode_models():
    ordered = [
        _normalize_reading_mode_model_name(GEMINI_READING_MODE_MODEL),
        "gemini-3-flash-preview",
        "gemini-2.5-flash",
        "gemini-flash-latest",
    ]
    seen = set()
    for model_name in ordered:
        token = str(model_name or "").strip()
        if not token or token in seen:
            continue
        seen.add(token)
        yield token


def _model_supports_thinking_level(model_name: str) -> bool:
    token = str(model_name or "").strip().lower()
    if not token:
        return False
    # As of Feb 2026, reading mode's fallback flash models reject thinking_level,
    # while Gemini 3 flash preview accepts it.
    return token.startswith("gemini-3")


def _extract_live_message_text(message):
    text_chunks = []
    server_content = getattr(message, "server_content", None)
    model_turn = getattr(server_content, "model_turn", None) if server_content else None
    parts = getattr(model_turn, "parts", None) if model_turn else None
    if parts:
        for part in parts:
            part_text = getattr(part, "text", None)
            if part_text:
                text_chunks.append(str(part_text))

    # Native-audio live models emit text via output_transcription.
    output_transcription = getattr(server_content, "output_transcription", None) if server_content else None
    if output_transcription and getattr(output_transcription, "text", None):
        text_chunks.append(str(output_transcription.text))
    return "".join(text_chunks).strip()


def _extract_live_audio_bytes(message):
    server_content = getattr(message, "server_content", None)
    model_turn = getattr(server_content, "model_turn", None) if server_content else None
    parts = getattr(model_turn, "parts", None) if model_turn else None
    if not parts:
        return b"", ""

    audio_bytes = []
    audio_mime = ""
    for part in parts:
        inline_data = getattr(part, "inline_data", None)
        if not inline_data:
            continue
        mime_type = str(getattr(inline_data, "mime_type", "") or "")
        data = getattr(inline_data, "data", None)
        if not data:
            continue
        if isinstance(data, bytes):
            payload = data
        elif isinstance(data, bytearray):
            payload = bytes(data)
        elif isinstance(data, str):
            try:
                payload = base64.b64decode(data, validate=True)
            except Exception:
                continue
        else:
            continue
        if not payload:
            continue
        if mime_type.lower().startswith("audio/"):
            if not audio_mime:
                audio_mime = mime_type
            audio_bytes.append(payload)
    return (b"".join(audio_bytes), audio_mime)


def _extract_pcm_rate(mime_type: str, default_rate: int = 24000) -> int:
    token = str(mime_type or "").strip().lower()
    match = re.search(r"rate\s*=\s*(\d+)", token)
    if not match:
        return default_rate
    try:
        return max(8000, int(match.group(1)))
    except Exception:
        return default_rate


def _pcm16le_to_wav_bytes(pcm_bytes: bytes, sample_rate: int = 24000, channels: int = 1) -> bytes:
    if not pcm_bytes:
        return b""
    sample_width = 2
    data_size = len(pcm_bytes)
    byte_rate = sample_rate * channels * sample_width
    block_align = channels * sample_width
    file_size = 36 + data_size
    header = (
        b"RIFF"
        + file_size.to_bytes(4, "little")
        + b"WAVE"
        + b"fmt "
        + (16).to_bytes(4, "little")
        + (1).to_bytes(2, "little")
        + channels.to_bytes(2, "little")
        + sample_rate.to_bytes(4, "little")
        + byte_rate.to_bytes(4, "little")
        + block_align.to_bytes(2, "little")
        + (16).to_bytes(2, "little")
        + b"data"
        + data_size.to_bytes(4, "little")
    )
    return header + pcm_bytes


async def _run_live_query_once(
    client,
    model_name: str,
    *,
    audio_bytes: bytes,
    audio_mime_type: str,
    screenshot_bytes: bytes,
    screenshot_mime_type: str,
    page_url: str,
    resume_handle: str = "",
    send_screenshot: bool = True,
):
    normalized_model_name = _normalize_live_model_name(model_name)
    connect_started_at = time.perf_counter()

    config = types.LiveConnectConfig(
        # Native-audio models are served as audio turns; recover text via output transcription.
        response_modalities=["AUDIO"],
        temperature=GEMINI_LIVE_TEMPERATURE,
        max_output_tokens=GEMINI_LIVE_MAX_OUTPUT_TOKENS,
        thinking_config=types.ThinkingConfig(thinking_budget=GEMINI_LIVE_THINKING_BUDGET),
        media_resolution=types.MediaResolution.MEDIA_RESOLUTION_LOW,
        realtime_input_config=types.RealtimeInputConfig(
            automatic_activity_detection=types.AutomaticActivityDetection(disabled=True),
        ),
        input_audio_transcription=types.AudioTranscriptionConfig(),
        output_audio_transcription=types.AudioTranscriptionConfig(),
        system_instruction=(
            "You are AQual, an accessibility assistant. "
            "Respond naturally in English. "
            "Do not include internal reasoning, analysis steps, or headings. "
            "Keep spoken replies short: one sentence when possible, at most two."
        ),
        session_resumption=types.SessionResumptionConfig(
            handle=resume_handle.strip() if resume_handle else None,
        ),
    )

    async with client.aio.live.connect(model=normalized_model_name, config=config) as session:
        connected_ms = round((time.perf_counter() - connect_started_at) * 1000, 1)
        send_started_at = time.perf_counter()
        await session.send_realtime_input(activity_start=types.ActivityStart())
        if send_screenshot and screenshot_bytes:
            await session.send_realtime_input(
                media=types.Blob(data=screenshot_bytes, mime_type=screenshot_mime_type),
            )
        if page_url:
            await session.send_realtime_input(text=f"Current page URL: {page_url}")
        await session.send_realtime_input(
            audio=types.Blob(data=audio_bytes, mime_type=audio_mime_type),
        )
        await session.send_realtime_input(activity_end=types.ActivityEnd())
        input_sent_ms = round((time.perf_counter() - send_started_at) * 1000, 1)

        answer_chunks = []
        transcript_text = ""
        output_transcript_text = ""
        output_audio_chunks = []
        output_audio_mime = ""
        next_session_handle = ""
        first_response_ms = 0.0
        receive_started_at = time.perf_counter()
        async for message in session.receive():
            chunk = _extract_live_message_text(message)
            if chunk:
                answer_chunks.append(chunk)
                if not first_response_ms:
                    first_response_ms = round((time.perf_counter() - receive_started_at) * 1000, 1)

            audio_chunk, audio_mime = _extract_live_audio_bytes(message)
            if audio_chunk:
                output_audio_chunks.append(audio_chunk)
                if not output_audio_mime:
                    output_audio_mime = audio_mime
                if not first_response_ms:
                    first_response_ms = round((time.perf_counter() - receive_started_at) * 1000, 1)

            server_content = getattr(message, "server_content", None)
            input_transcription = getattr(server_content, "input_transcription", None) if server_content else None
            if input_transcription and getattr(input_transcription, "text", None):
                transcript_text = str(input_transcription.text).strip()
            output_transcription = getattr(server_content, "output_transcription", None) if server_content else None
            if output_transcription and getattr(output_transcription, "text", None):
                output_transcript_text = str(output_transcription.text).strip()
            session_update = getattr(message, "session_resumption_update", None)
            if session_update and getattr(session_update, "resumable", None):
                new_handle = getattr(session_update, "new_handle", None)
                if new_handle:
                    next_session_handle = str(new_handle).strip()

            if server_content and getattr(server_content, "turn_complete", False):
                break

        answer = "".join(answer_chunks).strip()
        if not answer:
            answer = output_transcript_text
        if transcript_text and not re.search(r"[A-Za-z]", transcript_text):
            answer = "I could not clearly hear an English question. Please hold Alt+D and try again in English."
        output_audio_pcm = b"".join(output_audio_chunks)
        output_audio_wav = b""
        if output_audio_pcm:
            pcm_rate = _extract_pcm_rate(output_audio_mime, 24000)
            output_audio_wav = _pcm16le_to_wav_bytes(output_audio_pcm, sample_rate=pcm_rate, channels=1)
        if not answer and output_audio_wav:
            answer = "Spoken response generated."
        return {
            "answer": answer,
            "transcript": transcript_text,
            "model": model_name,
            "session_handle": next_session_handle,
            "audio_base64": base64.b64encode(output_audio_wav).decode("ascii") if output_audio_wav else "",
            "audio_mime": "audio/wav" if output_audio_wav else "",
            "audio_bytes": len(output_audio_wav),
            "connected_ms": connected_ms,
            "input_sent_ms": input_sent_ms,
            "first_response_ms": first_response_ms,
        }


async def _run_live_query_with_fallbacks(
    client,
    *,
    audio_bytes: bytes,
    audio_mime_type: str,
    screenshot_bytes: bytes,
    screenshot_mime_type: str,
    page_url: str,
    resume_handle: str = "",
    send_screenshot: bool = True,
):
    attempts = []
    for model_name in _iter_live_models():
        started_at = time.perf_counter()
        try:
            result = await _run_live_query_once(
                client,
                model_name,
                audio_bytes=audio_bytes,
                audio_mime_type=audio_mime_type,
                screenshot_bytes=screenshot_bytes,
                screenshot_mime_type=screenshot_mime_type,
                page_url=page_url,
                resume_handle=resume_handle,
                send_screenshot=send_screenshot,
            )
            duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
            _log_gemini_event(
                "live_model_success",
                request_kind="gemini_live_page_query",
                model=model_name,
                duration_ms=duration_ms,
                resumed_session=bool(resume_handle),
                screenshot_sent=bool(send_screenshot),
                connected_ms=float(result.get("connected_ms", 0.0) or 0.0),
                input_sent_ms=float(result.get("input_sent_ms", 0.0) or 0.0),
                first_response_ms=float(result.get("first_response_ms", 0.0) or 0.0),
                thinking_budget=GEMINI_LIVE_THINKING_BUDGET,
                max_output_tokens=GEMINI_LIVE_MAX_OUTPUT_TOKENS,
                temperature=GEMINI_LIVE_TEMPERATURE,
                answer_chars=len(result.get("answer", "")),
                transcript_chars=len(result.get("transcript", "")),
                output_audio_bytes=int(result.get("audio_bytes", 0) or 0),
            )
            if result.get("answer"):
                return result
            attempts.append(f"{model_name}: empty response")
        except Exception as e:
            duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
            attempts.append(f"{model_name}: {type(e).__name__}: {str(e)[:220]}")
            _log_gemini_event(
                "live_model_error",
                request_kind="gemini_live_page_query",
                model=model_name,
                duration_ms=duration_ms,
                error_type=type(e).__name__,
                error_message=str(e)[:240],
            )

    raise RuntimeError("; ".join(attempts) if attempts else "No Gemini Live model is configured")


def _generate_ai_json(
    client,
    parts,
    request_kind="generic",
    log_context=None,
    *,
    model_name: str = "",
    thinking_level: str = "",
    max_output_tokens: int = 0,
    temperature: float = None,
    response_mime_type: str = "",
    thinking_enabled: bool = True,
    use_stream: bool = True,
    response_schema=None,
    response_json_schema=None,
):
    selected_model = str(model_name or GEMINI_MODEL).strip() or GEMINI_MODEL
    selected_thinking_level = str(thinking_level or GEMINI_THINKING_LEVEL).strip() or GEMINI_THINKING_LEVEL
    request_id = uuid.uuid4().hex[:10]
    started_at = time.perf_counter()
    context = dict(log_context or {})
    _log_gemini_event(
        "request_start",
        request_id=request_id,
        request_kind=request_kind,
        model=selected_model,
        thinking_level=selected_thinking_level if thinking_enabled else "",
        max_output_tokens=int(max_output_tokens or 0),
        temperature=temperature if temperature is not None else "",
        response_mime_type=str(response_mime_type or "").strip(),
        use_stream=bool(use_stream),
        has_response_schema=bool(response_schema or response_json_schema),
        **context,
    )

    contents = [
        types.Content(
            role="user",
            parts=parts,
        ),
    ]
    config_kwargs = {}
    if thinking_enabled:
        config_kwargs["thinking_config"] = types.ThinkingConfig(
            thinking_level=selected_thinking_level,
        )
    if max_output_tokens and int(max_output_tokens) > 0:
        config_kwargs["max_output_tokens"] = int(max_output_tokens)
    if temperature is not None:
        config_kwargs["temperature"] = float(temperature)
    if response_mime_type:
        config_kwargs["response_mime_type"] = str(response_mime_type).strip()
    if response_schema is not None:
        config_kwargs["response_schema"] = response_schema
    if response_json_schema is not None:
        config_kwargs["response_json_schema"] = response_json_schema
    generate_content_config = types.GenerateContentConfig(**config_kwargs)

    response_text = ""
    chunk_count = 0
    try:
        if use_stream:
            for chunk in client.models.generate_content_stream(
                model=selected_model,
                contents=contents,
                config=generate_content_config,
            ):
                if chunk.text:
                    response_text += chunk.text
                    chunk_count += 1

            response_text = _clean_ai_json_text(response_text)
            try:
                parsed = json.loads(response_text)
            except json.JSONDecodeError:
                extracted = _extract_first_json_object_text(response_text)
                if not extracted:
                    raise
                parsed = json.loads(extracted)
        else:
            response = client.models.generate_content(
                model=selected_model,
                contents=contents,
                config=generate_content_config,
            )
            chunk_count = 1
            parsed = getattr(response, "parsed", None)
            if isinstance(parsed, str):
                parsed_text = _clean_ai_json_text(parsed)
                try:
                    parsed = json.loads(parsed_text)
                except json.JSONDecodeError:
                    extracted = _extract_first_json_object_text(parsed_text)
                    if not extracted:
                        raise
                    parsed = json.loads(extracted)
            if parsed is None:
                response_text = _clean_ai_json_text(getattr(response, "text", "") or "")
                try:
                    parsed = json.loads(response_text)
                except json.JSONDecodeError:
                    extracted = _extract_first_json_object_text(response_text)
                    if not extracted:
                        raise
                    parsed = json.loads(extracted)
            elif hasattr(parsed, "model_dump"):
                parsed = parsed.model_dump()
            if not response_text:
                try:
                    response_text = json.dumps(parsed, ensure_ascii=True, separators=(",", ":"))
                except Exception:
                    response_text = str(parsed)

        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_gemini_event(
            "request_success",
            request_id=request_id,
            request_kind=request_kind,
            duration_ms=duration_ms,
            response_chars=len(response_text),
            chunk_count=chunk_count,
            **context,
        )
        return parsed
    except Exception as e:
        if request_kind == "reading_mode_plan" and response_text:
            relaxed_payload = _parse_relaxed_reading_mode_payload(response_text)
            if relaxed_payload.get("include_selectors"):
                duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
                _log_gemini_event(
                    "request_success_relaxed",
                    request_id=request_id,
                    request_kind=request_kind,
                    duration_ms=duration_ms,
                    response_chars=len(response_text),
                    chunk_count=chunk_count,
                    **context,
                )
                return relaxed_payload
        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_gemini_event(
            "request_error",
            request_id=request_id,
            request_kind=request_kind,
            duration_ms=duration_ms,
            error_type=type(e).__name__,
            error_message=str(e)[:240],
            response_chars=len(response_text),
            response_preview=response_text[:180],
            chunk_count=chunk_count,
            **context,
        )
        raise


class BionicHtmlConverter:
    """Converts DOCX files to HTML with bionic reading formatting."""

    # XML namespaces used in DOCX
    NSMAP = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }

    def __init__(self, doc_id=None):
        self.css_classes = set()
        self.doc_id = doc_id
        self.image_counter = 0
        self.doc = None  # Will be set during convert

    def _get_fixation_length(self, word):
        """Calculates how many characters to bold for bionic reading."""
        clean_word = re.sub(r'\W+', '', word)
        length = len(clean_word)
        if length <= 1:
            return 1
        elif length <= 3:
            return int(length * 0.6) + 1
        else:
            return int(length * 0.45) + 1

    def _apply_bionic_to_text(self, text, is_already_bold=False):
        """Apply bionic reading formatting to text, returning HTML."""
        if is_already_bold:
            return f'<strong>{html.escape(text)}</strong>'

        parts = re.split(r'(\s+)', text)
        result = []

        for part in parts:
            if not part.strip() or not any(c.isalnum() for c in part):
                result.append(html.escape(part))
            else:
                fixation = self._get_fixation_length(part)
                bold_segment = html.escape(part[:fixation])
                normal_segment = html.escape(part[fixation:]) if len(part) > fixation else ''
                result.append(f'<strong class="bionic">{bold_segment}</strong>{normal_segment}')

        return ''.join(result)

    def _get_run_style(self, run):
        """Extract inline styles from a run."""
        styles = []

        if run.italic:
            styles.append('font-style: italic')
        if run.underline:
            styles.append('text-decoration: underline')
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            styles.append(f'color: #{rgb}')
        if run.font.size:
            pt_size = run.font.size.pt
            styles.append(f'font-size: {pt_size}pt')
        if run.font.name:
            styles.append(f'font-family: "{run.font.name}", sans-serif')

        return '; '.join(styles) if styles else None

    def _get_hyperlink_url(self, hyperlink_elem):
        """Extract URL from a hyperlink element."""
        r_id = hyperlink_elem.get(f'{{{self.NSMAP["r"]}}}id')
        if r_id and self.doc:
            try:
                rel = self.doc.part.rels.get(r_id)
                if rel and rel.target_ref:
                    return rel.target_ref
            except Exception:
                pass
        return None

    def _process_run(self, run, hyperlink_url=None):
        """Process a single run and return HTML."""
        text = run.text
        if not text:
            return ''

        style = self._get_run_style(run)
        style_attr = f' style="{style}"' if style else ''

        is_bold = run.bold is True
        bionic_html = self._apply_bionic_to_text(text, is_bold)

        if style:
            bionic_html = f'<span{style_attr}>{bionic_html}</span>'

        if hyperlink_url:
            return f'<a href="{html.escape(hyperlink_url)}" class="doc-link" target="_blank">{bionic_html}</a>'

        return bionic_html

    def _process_paragraph_content(self, paragraph):
        """Process paragraph content including hyperlinks."""
        content_parts = []
        p_elem = paragraph._p

        # Iterate through child elements to preserve order and detect hyperlinks
        for child in p_elem:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'r':  # Regular run
                # Find the corresponding run object
                for run in paragraph.runs:
                    if run._r == child:
                        content_parts.append(self._process_run(run))
                        break

            elif tag == 'hyperlink':  # Hyperlink
                url = self._get_hyperlink_url(child)
                # Process runs inside the hyperlink
                for r_elem in child.findall(f'{{{self.NSMAP["w"]}}}r'):
                    # Get text from the run
                    text_elems = r_elem.findall(f'{{{self.NSMAP["w"]}}}t')
                    for t_elem in text_elems:
                        if t_elem.text:
                            # Check for bold
                            rPr = r_elem.find(f'{{{self.NSMAP["w"]}}}rPr')
                            is_bold = False
                            if rPr is not None:
                                bold_elem = rPr.find(f'{{{self.NSMAP["w"]}}}b')
                                is_bold = bold_elem is not None

                            bionic_text = self._apply_bionic_to_text(t_elem.text, is_bold)
                            if url:
                                content_parts.append(f'<a href="{html.escape(url)}" class="doc-link" target="_blank">{bionic_text}</a>')
                            else:
                                content_parts.append(bionic_text)

        # Fallback if no content was extracted (handles simple paragraphs)
        if not content_parts:
            for run in paragraph.runs:
                content_parts.append(self._process_run(run))

        return ''.join(content_parts)

    def _get_list_info(self, paragraph):
        """Get list numbering info (type and level) from paragraph."""
        p_elem = paragraph._p
        numPr = p_elem.find(f'.//{{{self.NSMAP["w"]}}}numPr')

        if numPr is None:
            # Check style for list
            style_name = paragraph.style.name.lower() if paragraph.style else ''
            if 'list' in style_name:
                return {'is_list': True, 'level': 0, 'is_ordered': 'number' in style_name or 'decimal' in style_name}
            return None

        ilvl_elem = numPr.find(f'{{{self.NSMAP["w"]}}}ilvl')
        numId_elem = numPr.find(f'{{{self.NSMAP["w"]}}}numId')

        level = int(ilvl_elem.get(f'{{{self.NSMAP["w"]}}}val', 0)) if ilvl_elem is not None else 0
        num_id = numId_elem.get(f'{{{self.NSMAP["w"]}}}val') if numId_elem is not None else None

        # Try to determine if ordered or unordered from numbering definitions
        is_ordered = False
        if num_id and self.doc:
            try:
                numbering_part = self.doc.part.numbering_part
                if numbering_part:
                    # Check the abstract numbering for this numId
                    numbering_xml = numbering_part._element
                    for num in numbering_xml.findall(f'.//{{{self.NSMAP["w"]}}}num'):
                        if num.get(f'{{{self.NSMAP["w"]}}}numId') == num_id:
                            abstract_id = num.find(f'{{{self.NSMAP["w"]}}}abstractNumId')
                            if abstract_id is not None:
                                abs_id = abstract_id.get(f'{{{self.NSMAP["w"]}}}val')
                                for abstract in numbering_xml.findall(f'.//{{{self.NSMAP["w"]}}}abstractNum'):
                                    if abstract.get(f'{{{self.NSMAP["w"]}}}abstractNumId') == abs_id:
                                        for lvl in abstract.findall(f'{{{self.NSMAP["w"]}}}lvl'):
                                            if lvl.get(f'{{{self.NSMAP["w"]}}}ilvl') == str(level):
                                                numFmt = lvl.find(f'{{{self.NSMAP["w"]}}}numFmt')
                                                if numFmt is not None:
                                                    fmt = numFmt.get(f'{{{self.NSMAP["w"]}}}val', '')
                                                    is_ordered = fmt in ('decimal', 'lowerLetter', 'upperLetter', 'lowerRoman', 'upperRoman')
                                                break
                                        break
                            break
            except Exception:
                pass

        return {'is_list': True, 'level': level, 'is_ordered': is_ordered, 'num_id': num_id}

    def _get_paragraph_indent(self, paragraph):
        """Get paragraph indentation in ems."""
        p_elem = paragraph._p
        pPr = p_elem.find(f'{{{self.NSMAP["w"]}}}pPr')
        if pPr is None:
            return 0

        ind = pPr.find(f'{{{self.NSMAP["w"]}}}ind')
        if ind is None:
            return 0

        # Get left indentation in twips (1/20 of a point)
        left = ind.get(f'{{{self.NSMAP["w"]}}}left') or ind.get(f'{{{self.NSMAP["w"]}}}start') or '0'
        try:
            twips = int(left)
            # Convert to approximate ems (720 twips = 0.5 inch ≈ 2em)
            return round(twips / 360, 1)
        except ValueError:
            return 0

    def _get_paragraph_tag_and_class(self, paragraph):
        """Determine the appropriate HTML tag based on paragraph style."""
        style_name = paragraph.style.name.lower() if paragraph.style else ''

        if 'heading 1' in style_name:
            return 'h1', 'doc-heading doc-h1'
        elif 'heading 2' in style_name:
            return 'h2', 'doc-heading doc-h2'
        elif 'heading 3' in style_name:
            return 'h3', 'doc-heading doc-h3'
        elif 'heading 4' in style_name:
            return 'h4', 'doc-heading doc-h4'
        elif 'heading 5' in style_name:
            return 'h5', 'doc-heading doc-h5'
        elif 'heading 6' in style_name:
            return 'h6', 'doc-heading doc-h6'
        elif 'title' in style_name:
            return 'h1', 'doc-title'
        elif 'subtitle' in style_name:
            return 'h2', 'doc-subtitle'
        else:
            return 'p', 'doc-paragraph'

    def _get_alignment_style(self, paragraph):
        """Get CSS text-align from paragraph alignment."""
        alignment = paragraph.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return 'text-align: center'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return 'text-align: right'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            return 'text-align: justify'
        return None

    def _process_paragraph(self, paragraph, skip_list_check=False):
        """Process a paragraph and return HTML."""
        text = paragraph.text.strip()
        if not text:
            return '<p class="doc-paragraph">&nbsp;</p>'

        # Check for list
        if not skip_list_check:
            list_info = self._get_list_info(paragraph)
            if list_info:
                return None  # Will be handled by list processing

        tag, css_class = self._get_paragraph_tag_and_class(paragraph)
        alignment = self._get_alignment_style(paragraph)
        indent = self._get_paragraph_indent(paragraph)

        content = self._process_paragraph_content(paragraph)

        styles = []
        if alignment:
            styles.append(alignment)
        if indent > 0:
            styles.append(f'margin-left: {indent}em')

        style_attr = f' style="{"; ".join(styles)}"' if styles else ''

        return f'<{tag} class="{css_class}"{style_attr}>{content}</{tag}>'

    def _extract_images_from_paragraph(self, paragraph, doc):
        """Extract images from a paragraph and return HTML for them."""
        images_html = []

        # Look for drawings in the paragraph XML
        for run in paragraph.runs:
            run_xml = run._r
            drawings = run_xml.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

            for drawing in drawings:
                # Find the blip element which contains the image reference
                blips = drawing.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')

                for blip in blips:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        try:
                            image_part = doc.part.related_parts.get(embed_id)
                            if image_part:
                                image_bytes = image_part.blob
                                content_type = image_part.content_type

                                # Generate image ID
                                img_id = f"{self.doc_id}_{self.image_counter}"
                                self.image_counter += 1

                                # Store image data
                                image_store[img_id] = {
                                    'data': base64.b64encode(image_bytes).decode('utf-8'),
                                    'content_type': content_type
                                }

                                # Create base64 data URL for display
                                data_url = f"data:{content_type};base64,{base64.b64encode(image_bytes).decode('utf-8')}"

                                images_html.append(
                                    f'<div class="doc-image-container">'
                                    f'<img src="{data_url}" class="doc-image" data-image-id="{img_id}" alt="Document image">'
                                    f'<div class="image-description" style="display:none;"></div>'
                                    f'</div>'
                                )
                        except Exception:
                            pass

        return images_html

    def _process_table(self, table):
        """Process a table and return HTML."""
        rows_html = []

        for row in table.rows:
            cells_html = []
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        # For table cells, we process inline without the paragraph wrapper
                        parts = []
                        for run in paragraph.runs:
                            parts.append(self._process_run(run))
                        cell_content.append(''.join(parts))
                    else:
                        cell_content.append('&nbsp;')

                cells_html.append(f'<td class="doc-table-cell">{"<br>".join(cell_content)}</td>')

            rows_html.append(f'<tr>{"".join(cells_html)}</tr>')

        return f'<table class="doc-table"><tbody>{"".join(rows_html)}</tbody></table>'

    def convert(self, docx_file):
        """Convert a DOCX file to HTML with bionic reading formatting."""
        doc = Document(docx_file)
        self.doc = doc  # Store reference for helper methods
        html_parts = []

        # Track list state for proper nested list handling
        # Stack holds tuples of (is_ordered, level)
        list_stack = []
        current_list_items = []

        def close_lists_to_level(target_level):
            """Close nested lists down to target level."""
            nonlocal list_stack, current_list_items
            while list_stack and list_stack[-1][1] >= target_level:
                is_ordered, level = list_stack.pop()
                tag = 'ol' if is_ordered else 'ul'
                if current_list_items:
                    list_html = f'<{tag} class="doc-list doc-list-level-{level}">{"".join(current_list_items)}</{tag}>'
                    if list_stack:
                        # Nest inside parent list item
                        current_list_items = [list_html]
                    else:
                        html_parts.append(list_html)
                        current_list_items = []

        def close_all_lists():
            """Close all open lists."""
            nonlocal list_stack, current_list_items
            while list_stack:
                is_ordered, level = list_stack.pop()
                tag = 'ol' if is_ordered else 'ul'
                if current_list_items:
                    list_html = f'<{tag} class="doc-list doc-list-level-{level}">{"".join(current_list_items)}</{tag}>'
                    if list_stack:
                        current_list_items = [list_html]
                    else:
                        html_parts.append(list_html)
                        current_list_items = []

        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                for paragraph in doc.paragraphs:
                    if paragraph._p == element:
                        # Extract any images from this paragraph
                        images = self._extract_images_from_paragraph(paragraph, doc)
                        for img_html in images:
                            close_all_lists()
                            html_parts.append(img_html)

                        # Check for list formatting
                        list_info = self._get_list_info(paragraph)

                        if list_info:
                            level = list_info['level']
                            is_ordered = list_info['is_ordered']

                            # Handle list level changes
                            if not list_stack:
                                # Starting a new list
                                list_stack.append((is_ordered, level))
                                current_list_items = []
                            elif level > list_stack[-1][1]:
                                # Going deeper - start nested list
                                list_stack.append((is_ordered, level))
                            elif level < list_stack[-1][1]:
                                # Going up - close nested lists
                                while list_stack and list_stack[-1][1] > level:
                                    old_ordered, old_level = list_stack.pop()
                                    tag = 'ol' if old_ordered else 'ul'
                                    if current_list_items:
                                        nested_html = f'<{tag} class="doc-list doc-list-level-{old_level}">{"".join(current_list_items)}</{tag}>'
                                        current_list_items = [f'<li class="doc-list-item">{nested_html}</li>']

                                if not list_stack:
                                    list_stack.append((is_ordered, level))

                            # Process list item content using the hyperlink-aware method
                            if paragraph.text.strip():
                                content = self._process_paragraph_content(paragraph)
                                indent = self._get_paragraph_indent(paragraph)
                                style = f' style="margin-left: {indent}em"' if indent > 0 else ''
                                current_list_items.append(f'<li class="doc-list-item"{style}>{content}</li>')
                        else:
                            # Not a list item - close any open lists
                            close_all_lists()

                            if paragraph.text.strip():
                                para_html = self._process_paragraph(paragraph, skip_list_check=True)
                                if para_html:
                                    html_parts.append(para_html)
                        break

            # Check if it's a table
            elif element.tag.endswith('tbl'):
                close_all_lists()

                for table in doc.tables:
                    if table._tbl == element:
                        html_parts.append(self._process_table(table))
                        break

        # Close any remaining open lists
        close_all_lists()

        return '\n'.join(html_parts)


@app.route('/')
def index():
    """Serve the main page."""
    return render_template('index.html')


@app.route('/convert', methods=['POST'])
def convert_docx():
    """Handle DOCX file upload and conversion."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Only DOCX files are supported'}), 400

    try:
        # Read file into memory
        file_bytes = io.BytesIO(file.read())

        # Generate a unique document ID for image storage
        import uuid
        doc_id = str(uuid.uuid4())[:8]

        # Convert to HTML with bionic reading
        converter = BionicHtmlConverter(doc_id=doc_id)
        html_content = converter.convert(file_bytes)

        return jsonify({
            'success': True,
            'html': html_content,
            'filename': secure_filename(file.filename),
            'docId': doc_id
        })

    except Exception as e:
        return jsonify({'error': f'Failed to process document: {str(e)}'}), 500


@app.route('/summarize', methods=['POST'])
def summarize_paragraph():
    """Summarize a paragraph into bullet points using Gemini."""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400

    text = data['text'].strip()
    if not text:
        return jsonify({'error': 'Empty text'}), 400

    try:
        client = get_gemini_client()

        prompt = f"""Summarize the following paragraph into concise bullet points.
Return ONLY valid JSON in this exact format, nothing else:
{{"bullets": ["point 1", "point 2", "point 3"]}}

Paragraph:
{text}"""

        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part.from_text(text=prompt),
                ],
            ),
        ]

        generate_content_config = types.GenerateContentConfig(
            thinking_config=types.ThinkingConfig(
                thinking_level="LOW",
            ),
        )

        response_text = ""
        for chunk in client.models.generate_content_stream(
            model="gemini-3-flash-preview",
            contents=contents,
            config=generate_content_config,
        ):
            if chunk.text:
                response_text += chunk.text

        response_text = response_text.strip()

        # Clean up response - remove markdown code blocks if present
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            response_text = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])

        result = json.loads(response_text)
        return jsonify(result)

    except json.JSONDecodeError:
        return jsonify({'error': 'Failed to parse AI response'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/rephrase', methods=['POST'])
def rephrase_paragraph():
    """Rephrase a paragraph using Gemini, optionally matching a writing style."""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400

    text = data['text'].strip()
    if not text:
        return jsonify({'error': 'Empty text'}), 400

    writing_sample = data.get('writingSample', '').strip()

    try:
        client = get_gemini_client()

        if writing_sample:
            prompt = f"""You are a writing assistant. Rephrase the following paragraph to match the writing style, vocabulary level, and tone of the provided writing sample.

WRITING SAMPLE (match this style):
{writing_sample}

PARAGRAPH TO REPHRASE:
{text}

Return ONLY valid JSON in this exact format, nothing else:
{{
  "rephrased": "the rephrased paragraph here",
  "terms": [
    {{"word": "difficult word", "definition": "simple definition"}},
    {{"word": "another term", "definition": "its definition"}}
  ]
}}

The "terms" array should contain any words or phrases in your rephrased text that might be unfamiliar or technical to a general reader. Keep definitions concise (under 15 words)."""
        else:
            prompt = f"""You are a writing assistant. Rephrase the following paragraph to make it clearer and easier to understand while preserving the original meaning. Use simpler vocabulary where possible, but don't oversimplify technical concepts.

PARAGRAPH TO REPHRASE:
{text}

Return ONLY valid JSON in this exact format, nothing else:
{{
  "rephrased": "the rephrased paragraph here",
  "terms": [
    {{"word": "difficult word", "definition": "simple definition"}},
    {{"word": "another term", "definition": "its definition"}}
  ]
}}

The "terms" array should contain any words or phrases in your rephrased text that might be unfamiliar or technical to a general reader. Keep definitions concise (under 15 words). If there are no difficult terms, return an empty array."""

        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part.from_text(text=prompt),
                ],
            ),
        ]

        generate_content_config = types.GenerateContentConfig(
            thinking_config=types.ThinkingConfig(
                thinking_level="LOW",
            ),
        )

        response_text = ""
        for chunk in client.models.generate_content_stream(
            model="gemini-3-flash-preview",
            contents=contents,
            config=generate_content_config,
        ):
            if chunk.text:
                response_text += chunk.text

        response_text = response_text.strip()

        # Clean up response - remove markdown code blocks if present
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            response_text = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])

        result = json.loads(response_text)
        return jsonify(result)

    except json.JSONDecodeError:
        return jsonify({'error': 'Failed to parse AI response'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/extract-text', methods=['POST'])
def extract_text():
    """Extract text from uploaded DOCX file for writing sample."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Only DOCX files are supported'}), 400

    try:
        file_bytes = io.BytesIO(file.read())
        doc = Document(file_bytes)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Limit to first ~2000 words for context
        full_text = '\n\n'.join(text_parts)
        words = full_text.split()
        if len(words) > 2000:
            full_text = ' '.join(words[:2000]) + '...'

        return jsonify({
            'success': True,
            'text': full_text,
            'filename': secure_filename(file.filename)
        })

    except Exception as e:
        return jsonify({'error': f'Failed to extract text: {str(e)}'}), 500


@app.route('/describe-image', methods=['POST'])
def describe_image():
    """Generate an accessibility description for an image using Gemini."""
    data = request.get_json()
    if not data or 'imageId' not in data:
        return jsonify({'error': 'No image ID provided'}), 400

    image_id = data['imageId']
    if image_id not in image_store:
        return jsonify({'error': 'Image not found'}), 404

    image_data = image_store[image_id]

    try:
        client = get_gemini_client()
        image_bytes = base64.b64decode(image_data['data'])

        prompt = """Describe this image for a visually impaired person. Provide a clear, detailed description that captures:
1. The main subject or content of the image
2. Important visual details (colors, layout, text if any)
3. The context or purpose of the image in a document

Keep the description concise but informative (2-4 sentences). Return ONLY valid JSON in this exact format:
{"description": "Your description here"}"""

        result = _generate_ai_json(
            client,
            [
                types.Part.from_bytes(
                    data=image_bytes,
                    mime_type=image_data['content_type']
                ),
                types.Part.from_text(text=prompt),
            ],
            request_kind="bionic_doc_image_description",
            log_context={
                "route": "/describe-image",
                "source": "bionic_doc",
                "image_id": image_id,
                "image_bytes": len(image_bytes),
                "mime_type": image_data.get("content_type", ""),
                "prompt_chars": len(prompt),
            },
        )
        return jsonify(result)

    except json.JSONDecodeError:
        return jsonify({'error': 'Failed to parse AI response'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/describe-web-image', methods=['POST', 'OPTIONS'])
def describe_web_image():
    """Generate an accessibility description for an arbitrary webpage image."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    data = request.get_json() or {}
    raw_image_url = (data.get("imageUrl") or "").strip()

    try:
        image_bytes, content_type = _load_web_image_payload(data)
        client = get_gemini_client()

        if data.get("imageData"):
            input_mode = "base64"
        elif raw_image_url.startswith("data:"):
            input_mode = "data_url"
        elif raw_image_url:
            input_mode = "remote_url"
        else:
            input_mode = "unknown"

        alt_text = (data.get("altText") or "").strip()
        title_text = (data.get("titleText") or "").strip()
        nearby_text = (data.get("contextText") or "").strip()
        hint_parts = []
        if alt_text:
            hint_parts.append(f"ALT text: {alt_text[:180]}")
        if title_text and title_text.lower() != alt_text.lower():
            hint_parts.append(f"Title: {title_text[:140]}")
        if nearby_text:
            hint_parts.append(f"Nearby text: {nearby_text[:180]}")
        hint_block = " | ".join(hint_parts)

        prompt = """Describe this image for a visually impaired person. Provide a clear, detailed description that captures:
1. The main subject or content of the image
2. Important visual details (colors, layout, text if any)
3. The context or purpose of the image in a document

Keep the description concise but informative (2-4 sentences). Return ONLY valid JSON in this exact format:
{"description": "Your description here"}"""

        if hint_block:
            prompt += f"\n\nOptional on-page hints (can be wrong): {hint_block}"

        result = _generate_ai_json(
            client,
            [
                types.Part.from_bytes(
                    data=image_bytes,
                    mime_type=content_type,
                ),
                types.Part.from_text(text=prompt),
            ],
            request_kind="web_image_description",
            log_context={
                "route": "/describe-web-image",
                "source": "webpage",
                "input_mode": input_mode,
                "image_host": _extract_host(raw_image_url),
                "page_host": _extract_host(data.get("pageUrl", "")),
                "image_bytes": len(image_bytes),
                "mime_type": content_type,
                "prompt_chars": len(prompt),
                "hint_chars": len(hint_block),
                "prompt_style": "bionic_like",
            },
        )
        description = (result.get("description") or "").strip()
        if not description:
            return jsonify({'error': 'Failed to generate description'}), 500
        return jsonify({"description": description})

    except httpx.HTTPStatusError as e:
        return jsonify({'error': f'Failed to fetch image URL ({e.response.status_code})'}), 400
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except json.JSONDecodeError:
        return jsonify({'error': 'Failed to parse AI response'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/ask-web-image', methods=['POST', 'OPTIONS'])
def ask_web_image():
    """Answer follow-up questions about a webpage image."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    data = request.get_json() or {}
    raw_image_url = (data.get("imageUrl") or "").strip()
    question = (data.get("question") or "").strip()
    if not question:
        return jsonify({'error': 'No follow-up question provided'}), 400

    try:
        image_bytes, content_type = _load_web_image_payload(data)
        client = get_gemini_client()

        if data.get("imageData"):
            input_mode = "base64"
        elif raw_image_url.startswith("data:"):
            input_mode = "data_url"
        elif raw_image_url:
            input_mode = "remote_url"
        else:
            input_mode = "unknown"

        base_description = (data.get("description") or "").strip()
        history = data.get("history") or []
        history_lines = []
        for item in history[-8:]:
            if not isinstance(item, dict):
                continue
            role = (item.get("role") or "").strip().lower()
            content = (item.get("content") or "").strip()
            if role not in ("user", "assistant") or not content:
                continue
            history_lines.append(f"{role.upper()}: {content[:600]}")
        history_block = "\n".join(history_lines) if history_lines else "None."

        prompt = f"""You are helping a visually impaired user understand one image from a webpage.

Base description:
{base_description or "None."}

Conversation so far:
{history_block}

User follow-up question:
{question}

Answer the question directly and clearly in 1-3 short paragraphs.
If the question cannot be answered from the image, say so briefly.
Return ONLY valid JSON in this exact format:
{{"answer": "your answer"}}"""

        result = _generate_ai_json(
            client,
            [
                types.Part.from_bytes(
                    data=image_bytes,
                    mime_type=content_type,
                ),
                types.Part.from_text(text=prompt),
            ],
            request_kind="web_image_followup",
            log_context={
                "route": "/ask-web-image",
                "source": "webpage",
                "input_mode": input_mode,
                "image_host": _extract_host(raw_image_url),
                "page_host": _extract_host(data.get("pageUrl", "")),
                "image_bytes": len(image_bytes),
                "mime_type": content_type,
                "prompt_chars": len(prompt),
                "question_chars": len(question),
                "history_items": len(history_lines),
            },
        )
        answer = (result.get("answer") or "").strip()
        if not answer:
            return jsonify({'error': 'Failed to generate answer'}), 500
        return jsonify({"answer": answer})

    except httpx.HTTPStatusError as e:
        return jsonify({'error': f'Failed to fetch image URL ({e.response.status_code})'}), 400
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except json.JSONDecodeError:
        return jsonify({'error': 'Failed to parse AI response'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/gemini-live-query', methods=['POST', 'OPTIONS'])
def gemini_live_query():
    """Process spoken query + screenshot context using Gemini Live."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    data = request.get_json() or {}
    request_id = uuid.uuid4().hex[:10]
    started_at = time.perf_counter()
    page_url = (data.get("pageUrl") or "").strip()
    conversation_id = _normalize_live_conversation_id(data.get("conversationId"))
    resume_handle = _get_live_session_handle(conversation_id) if conversation_id else ""

    try:
        audio_bytes, audio_mime_type = _load_live_audio_payload(data)
        screenshot_bytes = b""
        screenshot_mime_type = "image/jpeg"
        screenshot_hash = ""
        send_screenshot = False
        if data.get("screenshotDataUrl") or data.get("screenshotData"):
            screenshot_bytes, screenshot_mime_type = _load_live_screenshot_payload(data)
            screenshot_hash = _get_screenshot_hash(screenshot_bytes)
            send_screenshot = _should_send_screenshot(conversation_id, screenshot_hash)
        elif not resume_handle:
            raise ValueError("No screenshot provided")
        client = get_gemini_client()

        _log_gemini_event(
            "request_start",
            request_id=request_id,
            request_kind="gemini_live_page_query",
            route="/gemini-live-query",
            source="webpage",
            model=GEMINI_LIVE_MODEL,
            fallback_models=[m for m in _iter_live_models() if m != GEMINI_LIVE_MODEL],
            thinking_budget=GEMINI_LIVE_THINKING_BUDGET,
            max_output_tokens=GEMINI_LIVE_MAX_OUTPUT_TOKENS,
            temperature=GEMINI_LIVE_TEMPERATURE,
            page_host=_extract_host(page_url),
            conversation_id=conversation_id,
            resumed_session=bool(resume_handle),
            audio_bytes=len(audio_bytes),
            audio_mime=audio_mime_type,
            screenshot_bytes=len(screenshot_bytes),
            screenshot_mime=screenshot_mime_type,
            screenshot_hash=screenshot_hash,
            screenshot_sent=bool(send_screenshot),
        )

        result = asyncio.run(_run_live_query_with_fallbacks(
            client,
            audio_bytes=audio_bytes,
            audio_mime_type=audio_mime_type,
            screenshot_bytes=screenshot_bytes,
            screenshot_mime_type=screenshot_mime_type,
            page_url=page_url,
            resume_handle=resume_handle,
            send_screenshot=send_screenshot,
        ))
        answer = (result.get("answer") or "").strip()
        transcript = (result.get("transcript") or "").strip()
        if not answer and not (result.get("audio_base64") or "").strip():
            raise RuntimeError("Gemini Live returned an empty answer")
        next_session_handle = str(result.get("session_handle") or "").strip()
        if conversation_id and next_session_handle:
            _set_live_session_handle(conversation_id, next_session_handle)

        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_gemini_event(
            "request_success",
            request_id=request_id,
            request_kind="gemini_live_page_query",
            route="/gemini-live-query",
            source="webpage",
            duration_ms=duration_ms,
            model=result.get("model", ""),
            conversation_id=conversation_id,
            resumed_session=bool(resume_handle),
            next_session_handle=bool(next_session_handle),
            answer_chars=len(answer),
            transcript_chars=len(transcript),
            audio_bytes=len(audio_bytes),
            screenshot_bytes=len(screenshot_bytes),
            screenshot_sent=bool(send_screenshot),
            output_audio_bytes=int(result.get("audio_bytes", 0) or 0),
            connected_ms=float(result.get("connected_ms", 0.0) or 0.0),
            input_sent_ms=float(result.get("input_sent_ms", 0.0) or 0.0),
            first_response_ms=float(result.get("first_response_ms", 0.0) or 0.0),
        )

        return jsonify({
            "answer": answer,
            "transcript": transcript,
            "model": result.get("model", ""),
            "audioBase64": result.get("audio_base64", ""),
            "audioMimeType": result.get("audio_mime", ""),
            "debug": {
                "screenshotSent": bool(send_screenshot),
                "screenshotHash": screenshot_hash,
                "connectedMs": float(result.get("connected_ms", 0.0) or 0.0),
                "inputSentMs": float(result.get("input_sent_ms", 0.0) or 0.0),
                "firstResponseMs": float(result.get("first_response_ms", 0.0) or 0.0),
            },
        })

    except ValueError as e:
        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_gemini_event(
            "request_error",
            request_id=request_id,
            request_kind="gemini_live_page_query",
            route="/gemini-live-query",
            source="webpage",
            duration_ms=duration_ms,
            error_type=type(e).__name__,
            error_message=str(e)[:240],
        )
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_gemini_event(
            "request_error",
            request_id=request_id,
            request_kind="gemini_live_page_query",
            route="/gemini-live-query",
            source="webpage",
            duration_ms=duration_ms,
            error_type=type(e).__name__,
            error_message=str(e)[:240],
        )
        return jsonify({'error': str(e)}), 500


@app.route('/reading-mode-plan', methods=['POST', 'OPTIONS'])
def reading_mode_plan():
    """Generate CSS keep/remove selectors for AI-assisted reading mode."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    data = request.get_json(silent=True) or {}
    page_url = str(data.get("pageUrl") or "").strip()
    page_title = str(data.get("pageTitle") or "").strip()
    source_html = str(data.get("htmlSource") or "")
    original_html_chars = len(source_html)
    visible_text_preview = str(data.get("visibleTextPreview") or "").strip()
    force_ai_refinement = str(data.get("forceAiRefinement", "")).strip().lower() in ("1", "true", "yes", "on")
    ai_refinement_enabled = bool(force_ai_refinement or GEMINI_READING_MODE_ENABLE_AI_REFINEMENT)

    if not source_html:
        return jsonify({'error': 'No page HTML provided'}), 400

    source_html = source_html.replace("\x00", "")
    truncated_input = False
    if len(source_html) > GEMINI_READING_MODE_HTML_CHAR_LIMIT:
        source_html = source_html[:GEMINI_READING_MODE_HTML_CHAR_LIMIT]
        truncated_input = True

    _log_reading_mode_event(
        "request_received",
        route="/reading-mode-plan",
        page_host=_extract_host(page_url),
        page_title_chars=len(page_title),
        html_chars_original=original_html_chars,
        html_chars_sent=len(source_html),
        visible_text_chars=len(visible_text_preview),
        force_ai_refinement=bool(force_ai_refinement),
        ai_refinement_enabled=bool(ai_refinement_enabled),
        selector_cap_disabled=bool(READING_MODE_DISABLE_SELECTOR_CAP),
        max_include_selectors=int(READING_MODE_MAX_INCLUDE_SELECTORS),
        max_exclude_selectors=int(READING_MODE_MAX_EXCLUDE_SELECTORS),
        input_truncated=bool(truncated_input),
    )

    if ai_refinement_enabled:
        url_cached_payload = _get_cached_reading_mode_plan_by_url(page_url)
        if isinstance(url_cached_payload, dict) and url_cached_payload:
            cached_debug = url_cached_payload.get("debug") if isinstance(url_cached_payload.get("debug"), dict) else {}
            url_cached_payload["debug"] = {
                **cached_debug,
                "cacheHit": True,
                "urlCacheHit": True,
                "inputTruncated": bool(truncated_input),
                "htmlChars": len(source_html),
                "htmlCharsOriginal": original_html_chars,
                "aiRefinementEnabled": bool(ai_refinement_enabled),
                "forceAiRefinement": bool(force_ai_refinement),
            }
            _log_reading_mode_event(
                "plan_url_cache_hit",
                route="/reading-mode-plan",
                model=str(url_cached_payload.get("model") or ""),
                include_count=len(url_cached_payload.get("includeSelectors") or []),
                exclude_count=len(url_cached_payload.get("excludeSelectors") or []),
            )
            return jsonify(url_cached_payload)

    html_fingerprint = hashlib.sha1(source_html.encode("utf-8", errors="ignore")).hexdigest()
    cache_key_seed = (
        f"{page_url[:600]}|{page_title[:240]}|{len(source_html)}|{html_fingerprint}"
        f"|force_ai={1 if force_ai_refinement else 0}|ai_enabled={1 if ai_refinement_enabled else 0}"
    )
    cache_key = hashlib.sha1(cache_key_seed.encode("utf-8")).hexdigest()
    cached_payload = _get_cached_reading_mode_plan(cache_key)
    if isinstance(cached_payload, dict) and cached_payload:
        cached_debug = cached_payload.get("debug") if isinstance(cached_payload.get("debug"), dict) else {}
        cached_payload["debug"] = {
            **cached_debug,
            "cacheHit": True,
            "urlCacheHit": bool(cached_debug.get("urlCacheHit", False)),
            "inputTruncated": bool(truncated_input),
            "htmlChars": len(source_html),
            "htmlCharsOriginal": original_html_chars,
        }
        _log_reading_mode_event(
            "plan_cache_hit",
            route="/reading-mode-plan",
            model=str(cached_payload.get("model") or ""),
            include_count=len(cached_payload.get("includeSelectors") or []),
            exclude_count=len(cached_payload.get("excludeSelectors") or []),
        )
        return jsonify(cached_payload)

    try:
        heuristic_plan = _build_heuristic_reading_mode_plan(source_html)
        heuristic_include = _normalize_selector_list(
            heuristic_plan.get("include_selectors"),
            max_items=READING_MODE_MAX_INCLUDE_SELECTORS,
        )
        heuristic_exclude = _normalize_selector_list(
            heuristic_plan.get("exclude_selectors"),
            max_items=READING_MODE_MAX_EXCLUDE_SELECTORS,
        )
        heuristic_candidates = heuristic_plan.get("candidate_include") if isinstance(heuristic_plan.get("candidate_include"), list) else []

        _log_reading_mode_event(
            "heuristic_plan_ready",
            route="/reading-mode-plan",
            include_count=len(heuristic_include),
            exclude_count=len(heuristic_exclude),
            candidate_count=len(heuristic_candidates),
            heuristic_error=str(heuristic_plan.get("error") or "")[:200],
        )

        if not ai_refinement_enabled:
            _log_reading_mode_event(
                "ai_refinement_skipped",
                route="/reading-mode-plan",
                reason="disabled",
            )
            response_payload = {
                "planVersion": READING_MODE_PLAN_CACHE_VERSION,
                "includeSelectors": heuristic_include,
                "excludeSelectors": heuristic_exclude,
                "notes": "Applied heuristic reading mode plan.",
                "model": "heuristic-only",
                "debug": {
                    "inputTruncated": bool(truncated_input),
                    "htmlChars": len(source_html),
                    "htmlCharsOriginal": original_html_chars,
                    "fallbackUsed": False,
                    "cacheHit": False,
                    "urlCacheHit": False,
                    "aiRefinementEnabled": bool(ai_refinement_enabled),
                    "forceAiRefinement": bool(force_ai_refinement),
                    "modelErrors": [],
                },
            }
            _set_cached_reading_mode_plan(cache_key, response_payload)
            _log_reading_mode_event(
                "plan_ready",
                route="/reading-mode-plan",
                model="heuristic-only",
                fallback_used=False,
                include_count=len(heuristic_include),
                exclude_count=len(heuristic_exclude),
                notes_chars=len(response_payload["notes"]),
                model_error_count=0,
                input_truncated=bool(truncated_input),
            )
            return jsonify(response_payload)

        selected_model = ""
        model_errors = []
        result = {}
        client = None
        if GEMINI_API_KEY:
            try:
                client = get_gemini_client()
            except Exception as e:
                model_errors.append(f"client_init: {type(e).__name__}: {str(e)[:180]}")
                _log_reading_mode_event(
                    "model_client_error",
                    route="/reading-mode-plan",
                    error_type=type(e).__name__,
                    error_message=str(e)[:240],
                )
        else:
            _log_reading_mode_event(
                "ai_skipped",
                route="/reading-mode-plan",
                reason="missing_api_key",
            )

        if client is not None:
            if READING_MODE_DISABLE_SELECTOR_CAP:
                exclude_selector_limit_rule = "- exclude_selectors can contain as many selectors as needed for accurate clutter removal."
            else:
                exclude_selector_limit_rule = f"- exclude_selectors must contain 0 to {READING_MODE_MAX_EXCLUDE_SELECTORS} selectors for elements to hide inside kept containers."
            prompt = f"""Return only a valid JSON object with this exact schema:
{{
  "include_selectors": ["..."],
  "exclude_selectors": ["..."],
  "notes": "..."
}}

Rules:
- Output JSON only. No markdown, no code fences, no explanation.
- include_selectors must contain 1 to 6 selectors for page containers to keep.
{exclude_selector_limit_rule}
- include_selectors must target broad readable containers (main article body), not tiny leaf nodes.
- exclude_selectors must target clutter INSIDE kept containers only: related cards, promo blocks, tag chips, recommendation rails, newsletter boxes, share/social modules, ad/sponsored blocks.
- Exclude topic/tag/taxonomy clusters and inline recommendation modules inside the article body.
- Do not enumerate individual nav links, menu items, or one-off link selectors. Return concise reusable module selectors only.
- Never use broad exclude selectors like: "div", "section", "article", "main", "p", "span", "*", "main *", "article *".
- Never exclude core readable text containers or headings.
- Never use these selectors: "*", "html", "body", ":root".
- Do not include script/style/meta/link selectors.
- Use selectors that are present in the provided page source.
- Prefer stable selectors using id/class/attributes over volatile hashed classes.
- If uncertain, keep the element (be conservative).

Page URL: {page_url[:400]}
Page title: {page_title[:400]}
Visible text preview:
{visible_text_preview[:5000] or "None."}

Page source (HTML):
```html
{source_html}
```
"""

            request_context = {
                "route": "/reading-mode-plan",
                "source": "webpage",
                "page_host": _extract_host(page_url),
                "page_title_chars": len(page_title),
                "html_chars": len(source_html),
                "html_chars_original": original_html_chars,
                "visible_text_chars": len(visible_text_preview),
                "input_truncated": bool(truncated_input),
            }

            candidate_models = list(_iter_reading_mode_models())
            if force_ai_refinement:
                # Forced refinement should still be responsive; try primary model only.
                primary_model = candidate_models[0] if candidate_models else _normalize_reading_mode_model_name(GEMINI_READING_MODE_MODEL)
                candidate_models = [primary_model] if primary_model else []

            for candidate_model in candidate_models:
                use_thinking = GEMINI_READING_MODE_USE_THINKING and _model_supports_thinking_level(candidate_model)
                try:
                    _log_reading_mode_event(
                        "model_attempt",
                        route="/reading-mode-plan",
                        model=candidate_model,
                        thinking_enabled=bool(use_thinking),
                        thinking_level=GEMINI_READING_MODE_THINKING_LEVEL if use_thinking else "",
                    )
                    result = _generate_ai_json(
                        client,
                        [types.Part.from_text(text=prompt)],
                        request_kind="reading_mode_plan",
                        log_context=request_context,
                        model_name=candidate_model,
                        thinking_level=GEMINI_READING_MODE_THINKING_LEVEL if use_thinking else "",
                        max_output_tokens=GEMINI_READING_MODE_MAX_OUTPUT_TOKENS,
                        temperature=0.0,
                        response_mime_type="application/json",
                        thinking_enabled=use_thinking,
                        use_stream=False,
                    )
                    selected_model = candidate_model
                    break
                except Exception as e:
                    model_errors.append(f"{candidate_model}: {type(e).__name__}: {str(e)[:180]}")
                    _log_reading_mode_event(
                        "model_attempt_error",
                        route="/reading-mode-plan",
                        model=candidate_model,
                        error_type=type(e).__name__,
                        error_message=str(e)[:240],
                    )

        ai_include = _normalize_selector_list(
            result.get("include_selectors") or result.get("includeSelectors"),
            max_items=READING_MODE_MAX_INCLUDE_SELECTORS,
        ) if isinstance(result, dict) else []
        ai_exclude = _normalize_selector_list(
            result.get("exclude_selectors") or result.get("excludeSelectors"),
            max_items=READING_MODE_MAX_EXCLUDE_SELECTORS,
        ) if isinstance(result, dict) else []
        ai_exclude = _normalize_ai_exclude_selector_list(ai_exclude, max_items=READING_MODE_MAX_EXCLUDE_SELECTORS)
        notes = str((result or {}).get("notes") or "").strip()[:400] if isinstance(result, dict) else ""

        include_selectors = ai_include or heuristic_include
        exclude_selectors = _merge_reading_mode_exclude_selectors(
            ai_exclude,
            heuristic_exclude,
            max_items=READING_MODE_MAX_EXCLUDE_SELECTORS,
        )
        fallback_used = not bool(ai_include)
        if fallback_used and not notes:
            notes = "Applied heuristic reading mode plan."

        model_token = selected_model or "heuristic-fallback"
        ai_clutter_selector_count = sum(
            1
            for selector in ai_exclude
            if any(hint in str(selector or "").lower() for hint in READING_MODE_CLUTTER_PRIORITY_HINTS)
        )
        final_clutter_selector_count = sum(
            1
            for selector in exclude_selectors
            if any(hint in str(selector or "").lower() for hint in READING_MODE_CLUTTER_PRIORITY_HINTS)
        )
        _log_reading_mode_event(
            "plan_ready",
            route="/reading-mode-plan",
            model=model_token,
            fallback_used=bool(fallback_used),
            include_count=len(include_selectors),
            exclude_count=len(exclude_selectors),
            ai_include_count=len(ai_include),
            ai_exclude_count=len(ai_exclude),
            heuristic_exclude_count=len(heuristic_exclude),
            ai_clutter_selector_count=int(ai_clutter_selector_count),
            final_clutter_selector_count=int(final_clutter_selector_count),
            selector_cap_disabled=bool(READING_MODE_DISABLE_SELECTOR_CAP),
            max_include_selectors=int(READING_MODE_MAX_INCLUDE_SELECTORS),
            max_exclude_selectors=int(READING_MODE_MAX_EXCLUDE_SELECTORS),
            notes_chars=len(notes),
            model_error_count=len(model_errors),
            input_truncated=bool(truncated_input),
        )

        response_payload = {
            "planVersion": READING_MODE_PLAN_CACHE_VERSION,
            "includeSelectors": include_selectors,
            "excludeSelectors": exclude_selectors,
            "notes": notes,
            "model": model_token,
            "debug": {
                "inputTruncated": bool(truncated_input),
                "htmlChars": len(source_html),
                "htmlCharsOriginal": original_html_chars,
                "fallbackUsed": bool(fallback_used),
                "cacheHit": False,
                "urlCacheHit": False,
                "aiRefinementEnabled": bool(ai_refinement_enabled),
                "forceAiRefinement": bool(force_ai_refinement),
                "modelErrors": model_errors[:6],
            },
        }
        _set_cached_reading_mode_plan(cache_key, response_payload)
        model_token_lower = str(model_token or "").lower()
        if model_token_lower.startswith("gemini"):
            _set_cached_reading_mode_plan_by_url(page_url, response_payload)
        return jsonify(response_payload)

    except Exception as e:
        _log_reading_mode_event(
            "plan_error",
            route="/reading-mode-plan",
            error_type=type(e).__name__,
            error_message=str(e)[:240],
        )
        fallback_include = ["main article", "article", "main", "[role='main']", "#content", ".article"]
        fallback_exclude = ["nav", "header", "footer", "aside", "[class*='ad-' i]", "[class*='cookie' i]", "[class*='consent' i]"]
        response_payload = {
            "planVersion": READING_MODE_PLAN_CACHE_VERSION,
            "includeSelectors": fallback_include,
            "excludeSelectors": fallback_exclude,
            "notes": "Applied fallback reading mode plan.",
            "model": "emergency-fallback",
            "debug": {
                "inputTruncated": bool(truncated_input),
                "htmlChars": len(source_html),
                "htmlCharsOriginal": original_html_chars,
                "fallbackUsed": True,
                "cacheHit": False,
                "urlCacheHit": False,
                "aiRefinementEnabled": bool(ai_refinement_enabled),
                "forceAiRefinement": bool(force_ai_refinement),
                "routeError": f"{type(e).__name__}: {str(e)[:220]}",
            },
        }
        _set_cached_reading_mode_plan(cache_key, response_payload)
        return jsonify(response_payload)


@app.route('/ring-event/push', methods=['POST', 'OPTIONS'])
def ring_event_push():
    """Receive ring button events from a native/local monitor process."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    data = request.get_json(silent=True) or {}
    source = str(data.get("source") or "ring-monitor")
    cursor, ts = _register_ring_event(source=source, payload=data)
    return jsonify({
        "ok": True,
        "cursor": int(cursor),
        "lastEventTs": float(ts),
    })


@app.route('/ring-event/poll', methods=['GET', 'OPTIONS'])
def ring_event_poll():
    """Poll ring events using a monotonically increasing cursor."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    raw_cursor = request.args.get("cursor", "0")
    try:
        client_cursor = max(0, int(raw_cursor))
    except Exception:
        client_cursor = 0

    with ring_event_lock:
        server_cursor = int(ring_event_counter)
        last_ts = float(ring_event_last_ts)
        events = [
            event
            for event in ring_event_history
            if int(event.get("cursor", 0)) > client_cursor
        ]

    delta = max(0, server_cursor - client_cursor)
    if RING_POLL_DEBUG:
        _log_ring_event(
            "event_poll",
            client_cursor=client_cursor,
            server_cursor=server_cursor,
            delta=delta,
            events_count=len(events),
        )
    return jsonify({
        "ok": True,
        "cursor": server_cursor,
        "delta": delta,
        "lastEventTs": last_ts,
        "events": events,
    })


@app.route('/speak-text', methods=['POST', 'OPTIONS'])
def speak_text():
    """Convert selected text to speech using ElevenLabs."""
    if request.method == 'OPTIONS':
        return _set_cors_headers(jsonify({'ok': True}))

    data = request.get_json() or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({'error': 'No text provided'}), 400

    if not ELEVENLABS_API_KEY:
        return jsonify({'error': 'ELEVENLABS_API_KEY is not configured. Set it in .env.'}), 500

    voice_id = (data.get("voiceId") or ELEVENLABS_TTS_VOICE_ID or "").strip()
    if not voice_id:
        return jsonify({'error': 'No ElevenLabs voice ID configured'}), 500

    text = text[:2500]
    request_id = uuid.uuid4().hex[:10]
    started_at = time.perf_counter()
    _log_tts_event(
        "request_start",
        request_id=request_id,
        route="/speak-text",
        voice_id=voice_id,
        model_id=ELEVENLABS_TTS_MODEL_ID,
        text_chars=len(text),
    )

    payload = {
        "text": text,
        "model_id": ELEVENLABS_TTS_MODEL_ID
    }
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{urllib.parse.quote(voice_id, safe='')}/stream"
    headers = {
        "xi-api-key": ELEVENLABS_API_KEY,
        "Accept": "audio/mpeg",
        "Content-Type": "application/json",
    }

    try:
        response = httpx.post(url, json=payload, headers=headers, timeout=45.0)
        if response.status_code >= 400:
            duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
            _log_tts_event(
                "request_error",
                request_id=request_id,
                route="/speak-text",
                duration_ms=duration_ms,
                status_code=response.status_code,
                body_preview=(response.text or "")[:200],
            )
            return jsonify({'error': f'ElevenLabs TTS failed ({response.status_code})'}), 502

        audio_bytes = response.content
        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_tts_event(
            "request_success",
            request_id=request_id,
            route="/speak-text",
            duration_ms=duration_ms,
            status_code=response.status_code,
            audio_bytes=len(audio_bytes),
        )
        return Response(
            audio_bytes,
            mimetype="audio/mpeg",
            headers={"Cache-Control": "no-store"},
        )

    except httpx.RequestError as e:
        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_tts_event(
            "request_error",
            request_id=request_id,
            route="/speak-text",
            duration_ms=duration_ms,
            error_type=type(e).__name__,
            error_message=str(e)[:240],
        )
        return jsonify({'error': 'Could not reach ElevenLabs TTS service'}), 502
    except Exception as e:
        duration_ms = round((time.perf_counter() - started_at) * 1000, 1)
        _log_tts_event(
            "request_error",
            request_id=request_id,
            route="/speak-text",
            duration_ms=duration_ms,
            error_type=type(e).__name__,
            error_message=str(e)[:240],
        )
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=False, use_reloader=False, port=8080)
