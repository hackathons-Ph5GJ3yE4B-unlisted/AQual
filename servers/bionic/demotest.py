from docx import Document
import re
import os

class BionicDocxConverter:
    def __init__(self, input_path, output_path):
        self.input_path = input_path
        self.output_path = output_path

    def _get_fixation_length(self, word):
        """Calculates how many characters to bold."""
        clean_word = re.sub(r'\W+', '', word)
        length = len(clean_word)
        if length <= 1: return 1
        elif length <= 3: return int(length * 0.6) + 1
        else: return int(length * 0.45) + 1

    def _copy_style(self, source_run, target_run):
        """Copies formatting from source to target run."""
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def _process_paragraph(self, paragraph):
        """Surgically replaces text in a paragraph while respecting existing bolding."""
        # Work on a copy of the runs list because we will be modifying the paragraph
        original_runs = list(paragraph.runs)
        if not original_runs:
            return

        # Clear the paragraph to rebuild it
        # Note: We use this method to keep the paragraph object/style intact
        p_element = paragraph._p
        for run in original_runs:
            p_element.remove(run._r)

        for run in original_runs:
            # CHECK: If the run is already bold, do not apply bionic logic
            # run.bold can be True, False, or None (inherits from style)
            if run.bold is True:
                new_run = paragraph.add_run(run.text)
                self._copy_style(run, new_run)
                new_run.bold = True
                continue
            
            # Apply bionic reading to non-bold runs
            # Split by whitespace but keep the whitespace in the list
            parts = re.split(r'(\s+)', run.text)
            
            for part in parts:
                # If it's just whitespace or purely punctuation, add it as a normal run
                if not part.strip() or not any(c.isalnum() for c in part):
                    new_run = paragraph.add_run(part)
                    self._copy_style(run, new_run)
                    new_run.bold = run.bold # Maintain whatever it was (False/None)
                else:
                    # Bionic logic for words
                    fixation = self._get_fixation_length(part)
                    
                    # Bold part
                    bold_segment = part[:fixation]
                    b_run = paragraph.add_run(bold_segment)
                    self._copy_style(run, b_run)
                    b_run.bold = True
                    
                    # Normal part
                    if len(part) > fixation:
                        norm_segment = part[fixation:]
                        n_run = paragraph.add_run(norm_segment)
                        self._copy_style(run, n_run)
                        n_run.bold = False

    def convert(self):
        if not os.path.exists(self.input_path):
            print(f"Error: {self.input_path} not found.")
            return

        doc = Document(self.input_path)
        print(f"Processing '{self.input_path}'...")

        # Process standard paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                self._process_paragraph(paragraph)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            self._process_paragraph(paragraph)

        doc.save(self.output_path)
        print(f"Successfully saved Bionic version to: {self.output_path}")

if __name__ == "__main__":
    # Ensure you have 'input.docx' in the same directory
    tool = BionicDocxConverter("input.docx", "bionic_output.docx")
    tool.convert()